# promopro_v3.17.y.py
# - Клавиатура: 2 кнопки в ряд, согласованы подписи.
# - Обработчики: матч по подстроке (lower), не зависят от точного текста.

import asyncio
import datetime
import os
import json
import re
import logging
import traceback
from copy import deepcopy

from aiogram import Bot, Dispatcher, F
from aiogram.types import (
    Message, ReplyKeyboardMarkup, KeyboardButton,
    InlineKeyboardMarkup, InlineKeyboardButton, CallbackQuery, FSInputFile
)
from aiogram.filters import CommandStart, Command, StateFilter
from aiogram.fsm.state import StatesGroup, State
from aiogram.fsm.context import FSMContext
from aiogram.enums import ParseMode
from aiogram.client.default import DefaultBotProperties
from aiogram.client.session.aiohttp import AiohttpSession

try:
    from dotenv import load_dotenv
    load_dotenv()
except Exception:
    pass

# ── Импорт конфигурации ──
try:
    import config
    # Основные настройки
    BOT_TOKEN = config.BOT_TOKEN
    ADMIN_CHAT_ID = getattr(config, 'ADMIN_CHAT_ID', None)
    
    # VK.ОРД API настройки
    VK_ORD_API_TOKEN = getattr(config, 'VK_ORD_API_TOKEN', None)
    VK_ORD_API_BASE = getattr(config, 'VK_ORD_API_BASE', 'https://api-sandbox.ord.vk.com')
    VK_ORD_PERSON_TYPE_JURIDICAL = getattr(config, 'VK_ORD_PERSON_TYPE_JURIDICAL', 'juridical')
    VK_ORD_PERSON_TYPE_IP = getattr(config, 'VK_ORD_PERSON_TYPE_IP', 'ip')
    VK_ORD_PERSON_TYPE_INDIVIDUAL = getattr(config, 'VK_ORD_PERSON_TYPE_INDIVIDUAL', 'physical')
    VK_ORD_PERSON_TYPE_DEFAULT = getattr(config, 'VK_ORD_PERSON_TYPE_DEFAULT', 'juridical')
    TEMPLATE_INVOICE_SINGLE = getattr(config, 'TEMPLATE_INVOICE_SINGLE', 'templates/schet-oferta.docx')
    TEMPLATE_INVOICE_MULTI = getattr(config, 'TEMPLATE_INVOICE_MULTI', 'templates/schet-oferta2-multi.docx')
    TEMPLATE_INVOICE_MULTI_PRO = getattr(config, 'TEMPLATE_INVOICE_MULTI_PRO', 'templates/schet-oferta2-multiPRO.docx')
    TEMPLATE_CONTRACT = getattr(config, 'TEMPLATE_CONTRACT', 'templates/dogovor_rim.docx')
    TEMPLATE_CONTRACT_MULTI = getattr(config, 'TEMPLATE_CONTRACT_MULTI', 'templates/dogovor_rim2-multi.docx')
    OUTPUT_DIR = getattr(config, 'OUTPUT_DIR', 'generated')
    COUNTERS_FILE = getattr(config, 'COUNTERS_FILE', 'counters.json')
    METRICS_FILE = getattr(config, 'METRICS_FILE', 'metrics.json')
    MAX_ITEMS_FOR_TEMPLATE = getattr(config, 'MAX_ITEMS_FOR_TEMPLATE', 50)
    CAPTION_LIMIT = getattr(config, 'CAPTION_LIMIT', 1024)
except ImportError:
    raise SystemExit("Файл config.py не найден! Создайте его на основе config.example.py") from None

if not BOT_TOKEN:
    raise SystemExit("BOT_TOKEN не задан в config.py!")

try:
    from docx import Document
    from docx.shared import Pt
except Exception as e:
    raise SystemExit("Нужен python-docx: pip install python-docx") from e

# ── Логирование ──
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")

# ── TZ: Europe/Moscow ──
try:
    from zoneinfo import ZoneInfo
    TZ = ZoneInfo("Europe/Moscow")

    def now_tz() -> datetime.datetime:
        return datetime.datetime.now(TZ)
except Exception:

    def now_tz() -> datetime.datetime:
        return datetime.datetime.now()

# Создаем OUTPUT_DIR если не существует
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ================== FSM ====================
class InvoiceForm(StatesGroup):
    customer_name = State()
    customer_inn  = State()
    item_channel  = State()
    item_period   = State()
    item_amount   = State()
    manual_text   = State()
    manual_amount = State()
    confirm       = State()


class ContractForm(StatesGroup):
    customer_name      = State()
    customer_inn       = State()
    customer_ogrn      = State()
    placement_channel  = State()
    service_date       = State()
    service_period     = State()
    amount             = State()
    confirm            = State()

# ── Хелперы ──
def match_contains(substr: str):
    # Нормализуем пробелы и регистр для устойчивого матчинга по подстроке
    s = re.sub(r'\s+', ' ', str(substr).lower()).strip()

    def _pred(t):
        if not isinstance(t, str):
            return False
        norm = re.sub(r'\s+', ' ', t.lower()).strip()
        return s in norm

    return F.text.func(_pred)


def fmt_amount(n: int) -> str:
    return f"{n:,}".replace(",", " ")


MD_SAFE_PATTERN = re.compile(r'([_*`\[])')


def md_escape(s: str) -> str:
    if s is None:
        return ""
    return MD_SAFE_PATTERN.sub(r'\\\1', str(s))


def number_to_words_ru(n: int) -> str:
    """
    Корректное преобразование целого числа в строки с суммой прописью (RU),
    устойчивое для сумм вплоть до триллионов.
    """
    n = int(n)
    if n == 0:
        return "ноль"

    negative = n < 0
    if negative:
        n = -n

    units_male = ["", "один", "два", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    units_fem  = ["", "одна", "две", "три", "четыре", "пять", "шесть", "семь", "восемь", "девять"]
    teens = [
        "десять", "одиннадцать", "двенадцать", "тринадцать", "четырнадцать",
        "пятнадцать", "шестнадцать", "семнадцать", "восемнадцать", "девятнадцать",
    ]
    tens = [
        "", "", "двадцать", "тридцать", "сорок",
        "пятьдесят", "шестьдесят", "семьдесят", "восемьдесят", "девяносто",
    ]
    hundreds = [
        "", "сто", "двести", "триста", "четыреста",
        "пятьсот", "шестьсот", "семьсот", "восемьсот", "девятьсот",
    ]

    # (name_singular, name_few, name_many, female)
    groups = [
        ("", "", "", False),                   # единицы (рубли – слово добавляется снаружи)
        ("тысяча", "тысячи", "тысяч", True),   # тысячи
        ("миллион", "миллиона", "миллионов", False),
        ("миллиард", "миллиарда", "миллиардов", False),
        ("триллион", "триллиона", "триллионов", False),
    ]

    def tri(num: int, female: bool) -> str:
        assert 0 <= num <= 999
        words = []
        h = num // 100
        if h:
            words.append(hundreds[h])
        r = num % 100
        if 10 <= r <= 19:
            words.append(teens[r - 10])
        else:
            t = r // 10
            u = r % 10
            if t:
                words.append(tens[t])
            if u:
                words.append((units_fem if female else units_male)[u])
        return " ".join(words)

    parts = []
    group_index = 0
    while n > 0 and group_index < len(groups):
        num = n % 1000
        n //= 1000
        if num:
            name_s, name_few, name_many, female = groups[group_index]
            chunk_words = tri(num, female)
            # Выбор правильной формы слова группы
            if name_s:
                last_two = num % 100
                last = num % 10
                if 11 <= last_two <= 14:
                    name = name_many
                elif last == 1:
                    name = name_s
                elif last in (2, 3, 4):
                    name = name_few
                else:
                    name = name_many
                parts.append(f"{chunk_words} {name}".strip())
            else:
                parts.append(chunk_words)
        group_index += 1

    words = " ".join(reversed([p for p in parts if p])).strip()
    if negative:
        words = "минус " + words
    return words


DATE_RX = re.compile(r"\b(\d{1,2}\.\d{1,2}\.(?:\d{2}|\d{4}))\b")


def normalize_date_for_service_date(raw: str) -> str | None:
    m = DATE_RX.search(raw or "")
    if not m:
        return None
    d = m.group(1)
    if len(d) == 8:
        dd, mm, yy = d.split(".")
        return f"{dd}.{mm}.20{yy}"
    return d


def only_digits(n: str) -> int:
    digits = re.sub(r"[^\d]", "", n or "")
    return int(digits) if digits else 0

# ── Клавиатуры ──
def main_kb() -> ReplyKeyboardMarkup:
    return ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="💳 Выставить «Счёт на оплату»"),
             KeyboardButton(text="📃 Составить «Договор РИМ»"), ],
            [KeyboardButton(text="🔄 Сброс нумерации"),
             KeyboardButton(text="🔍 Поиск по ИНН")],
            [KeyboardButton(text="➦ Перейти в кабинет «VK.ОРД»")],
            [KeyboardButton(text="⚙️              Обратная связь                  ⚙️")]
        ],
        resize_keyboard=True,
        input_field_placeholder="Выберите действие…"
    )


def vk_lk_subscribe_kb() -> ReplyKeyboardMarkup:
    return ReplyKeyboardMarkup(
        keyboard=[
            [
                KeyboardButton(text="✔️ Да"),
                KeyboardButton(text="❌ Не надо"),
            ],
            [KeyboardButton(text="📚 Подробнее")],
        ],
        resize_keyboard=True
    )


def reset_confirm_kb() -> ReplyKeyboardMarkup:
    return ReplyKeyboardMarkup(
        keyboard=[
            [
                KeyboardButton(text="✔️ ДА"),
                KeyboardButton(text="❌ НЕТ"),
            ],
        ],
        resize_keyboard=True
    )


def step_kb() -> ReplyKeyboardMarkup:
    return ReplyKeyboardMarkup(
        keyboard=[[KeyboardButton(text="◀  Назад"), KeyboardButton(text="✖  На главную")]],
        resize_keyboard=True
    )


def invoice_actions_kb() -> ReplyKeyboardMarkup:
    return ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="➕  Добавить пункт в счёт")],
            [KeyboardButton(text="➕  Добавить вручную")],
            [KeyboardButton(text="✅  Сформировать счёт")],
            [KeyboardButton(text="◀  Назад"), KeyboardButton(text="✖  На главную")],
        ],
        resize_keyboard=True
    )


def contract_actions_kb() -> ReplyKeyboardMarkup:
    return ReplyKeyboardMarkup(
        keyboard=[
            [KeyboardButton(text="➕  Добавить пункт")],
            [KeyboardButton(text="✅  Сформировать договор")],
            [KeyboardButton(text="◀  Назад"), KeyboardButton(text="✖  На главную")],
        ],
        resize_keyboard=True
    )


def inline_new_invoice() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[[InlineKeyboardButton(text="➕  Создать ещё один счёт", callback_data="new_invoice")]]
    )


def inline_new_contract() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup(
        inline_keyboard=[[InlineKeyboardButton(text="🆕  Создать ещё один договор", callback_data="new_contract")]]
    )

# ── Счётчики ──
def load_counters() -> dict:
    if not os.path.exists(COUNTERS_FILE):
        return {}
    try:
        with open(COUNTERS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def save_counters(data: dict) -> None:
    with open(COUNTERS_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


# ── Метрики уникальных пользователей ──
def load_metrics() -> dict:
    """Загружает метрики уникальных пользователей из файла."""
    if not os.path.exists(METRICS_FILE):
        return {"unique_users": {}, "total_count": 0, "daily_registrations": {}}
    try:
        with open(METRICS_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
            # Инициализация структуры для обратной совместимости
            if "unique_users" not in data:
                data["unique_users"] = {}
            if "total_count" not in data:
                data["total_count"] = len(data.get("unique_users", {}))
            if "daily_registrations" not in data:
                data["daily_registrations"] = {}
            return data
    except Exception:
        return {"unique_users": {}, "total_count": 0, "daily_registrations": {}}


def save_metrics(data: dict) -> None:
    """Сохраняет метрики уникальных пользователей в файл."""
    with open(METRICS_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def track_unique_user(user_id: int) -> bool:
    """Отслеживает уникального пользователя. Возвращает True если пользователь новый."""
    metrics = load_metrics()
    user_id_str = str(user_id)
    
    if user_id_str not in metrics["unique_users"]:
        # Новый пользователь
        now = now_tz()
        date_key = now.strftime("%Y-%m-%d")
        metrics["unique_users"][user_id_str] = date_key
        metrics["total_count"] = len(metrics["unique_users"])
        
        # Увеличиваем счетчик регистраций за день
        if date_key not in metrics["daily_registrations"]:
            metrics["daily_registrations"][date_key] = 0
        metrics["daily_registrations"][date_key] += 1
        
        save_metrics(metrics)
        return True
    return False


def get_unique_users_count() -> int:
    """Возвращает общее количество уникальных пользователей."""
    metrics = load_metrics()
    return metrics.get("total_count", len(metrics.get("unique_users", {})))


def get_unique_users_stats() -> dict:
    """Возвращает статистику по уникальным пользователям."""
    metrics = load_metrics()
    now = now_tz()
    today_key = now.strftime("%Y-%m-%d")
    
    # Подсчет за период
    week_ago = now - datetime.timedelta(days=7)
    month_ago = now - datetime.timedelta(days=30)
    
    today_count = metrics["daily_registrations"].get(today_key, 0)
    
    week_count = 0
    month_count = 0
    for date_str, count in metrics["daily_registrations"].items():
        try:
            date = datetime.datetime.strptime(date_str, "%Y-%m-%d")
            if date >= week_ago:
                week_count += count
            if date >= month_ago:
                month_count += count
        except Exception:
            continue
    
    return {
        "total": metrics["total_count"],
        "today": today_count,
        "week": week_count,
        "month": month_count
    }


def get_user_daily_sequence(now: datetime.datetime, user_id: int) -> str:
    counters = load_counters()
    date_key = now.strftime("%Y-%m-%d")
    per_day = counters.get(date_key) or {}
    per_day[str(user_id)] = per_day.get(str(user_id), 0) + 1
    counters[date_key] = per_day
    save_counters(counters)
    return f"{per_day[str(user_id)]:02d}"


def reset_user_daily_sequence(now: datetime.datetime, user_id: int) -> int:
    counters = load_counters()
    date_key = now.strftime("%Y-%m-%d")
    per_day = counters.get(date_key) or {}
    prev = per_day.get(str(user_id), 0)
    per_day[str(user_id)] = 0
    counters[date_key] = per_day
    save_counters(counters)
    return prev


def generate_number(now: datetime.datetime, user_id: int) -> str:
    return f"{now.strftime('%d')}-{now.strftime('%m')}-{get_user_daily_sequence(now, user_id)}"


def generate_date(now: datetime.datetime) -> str:
    return now.strftime("%d.%m.%Y")

# ── Подстановка в DOCX ──
def _replace_in_paragraph(paragraph, mapping: dict):
    if not mapping:
        return
    text = paragraph.text or ""
    orig = text
    for k, v in mapping.items():
        if k in text:
            text = text.replace(k, v)
    if text != orig:
        if paragraph.runs:
            paragraph.runs[0].text = text
            for r in paragraph.runs[1:]:
                r.text = ""
        else:
            paragraph.add_run(text)


def _replace_in_table(table, mapping: dict):
    for row in table.rows:
        for cell in row.cells:
            _replace_in_block(cell, mapping)


def _replace_in_header_footer(hf, mapping: dict):
    for p in hf.paragraphs:
        _replace_in_paragraph(p, mapping)
    for t in hf.tables:
        _replace_in_table(t, mapping)


def _replace_in_block(container, mapping: dict):
    for p in getattr(container, "paragraphs", []):
        _replace_in_paragraph(p, mapping)
    for t in getattr(container, "tables", []):
        _replace_in_table(t, mapping)


TAG_CH = "{{PLACEMENT_CHANNEL}}"
TAG_PR = "{{SERVICE_PERIOD}}"
TAG_AM = "{{AMOUNT}}"
TAG_SD = "{{SERVICE_DATE}}"


def tag_n(base: str, i: int) -> str:
    if i == 1:
        return base
    m = re.fullmatch(r"\{\{([A-Z0-9_]+)\}\}", base)
    return "{{" + f"{m.group(1)}{i}" + "}}" if m else f"{base}{i}"


def enforce_times12_cell(cell):
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)


def copy_cell_alignment(src_cell, dst_cell):
    try:
        dst_cell.vertical_alignment = getattr(src_cell, "vertical_alignment", None)
    except Exception:
        pass
    if src_cell.paragraphs and dst_cell.paragraphs:
        sp = src_cell.paragraphs[0]
        dp = dst_cell.paragraphs[0]
        dp.alignment = sp.alignment
        pf_src, pf_dst = sp.paragraph_format, dp.paragraph_format
        pf_dst.left_indent = pf_src.left_indent
        pf_dst.right_indent = pf_src.right_indent
        pf_dst.first_line_indent = pf_src.first_line_indent
        pf_dst.space_before = pf_src.space_before
        pf_dst.space_after = pf_src.space_after
        pf_dst.line_spacing = pf_src.line_spacing


def replace_tag_in_cell(cell, old_tag: str, new_tag_text: str, tmpl_cell):
    new_text = (cell.text or "").replace(old_tag, new_tag_text or "")
    cell.text = new_text
    copy_cell_alignment(tmpl_cell, cell)
    enforce_times12_cell(cell)


def _cell_has(cell, needle: str) -> bool:
    try:
        return needle in (cell.text or "")
    except Exception:
        return False


def _find_invoice_table_and_columns(doc: Document):
    for tbl in doc.tables:
        row1 = row2 = None
        col_map = {'channel': None, 'period': None, 'amount': None, 'sdate': None}
        num_col_idx = None

        for ri, row in enumerate(tbl.rows):
            if any(_cell_has(c, TAG_CH) for c in row.cells):
                row1 = ri
            if any(_cell_has(c, tag_n(TAG_CH, 2)) for c in row.cells):
                row2 = ri
        if row1 is None or row2 is None:
            continue

        for ci, cell in enumerate(tbl.rows[row2].cells):
            t = cell.text
            if tag_n(TAG_CH, 2) in t:
                col_map['channel'] = ci
            if tag_n(TAG_PR, 2) in t:
                col_map['period'] = ci
            if tag_n(TAG_AM, 2) in t:
                col_map['amount'] = ci
            if tag_n(TAG_SD, 2) in t:
                col_map['sdate'] = ci

        def looks_like_num(s: str, n: int) -> bool:
            s = (s or "").strip()
            return s == str(n) or s == f"{n}." or s.startswith(f"{n} ")

        for ci, cell in enumerate(tbl.rows[row1].cells):
            if looks_like_num(cell.text, 1):
                num_col_idx = ci
                break
        if num_col_idx is None:
            num_col_idx = 0

        return tbl, row1, row2, col_map, num_col_idx
    return None, None, None, None, None


def _insert_template_row_after(table, template_row, after_row_index: int):
    anchor_tr = table.rows[after_row_index]._tr
    new_tr = deepcopy(template_row._tr)
    anchor_tr.addnext(new_tr)
    return table.rows[after_row_index + 1]


def _infer_number_format(row1_text: str, row2_text: str):
    if row1_text.strip().endswith(".") and row2_text.strip().endswith("."):
        return lambda n: f"{n}."
    if row1_text.strip().isdigit() and row2_text.strip().isdigit():
        return lambda n: f"{n}"
    m = re.search(r"(\D*)(\d+)(\D*)", row1_text or "")
    if m:
        prefix, _, suffix = m.groups()
        return lambda n: f"{prefix}{n}{suffix}"
    return lambda n: f"{n}."


def add_dynamic_rows_for_items(doc: Document, items: list):
    if len(items) <= 2:
        return
    tbl, r1, r2, col_map, num_col = _find_invoice_table_and_columns(doc)
    if tbl is None:
        return

    cell1_num = tbl.rows[r1].cells[num_col].text
    cell2_num = tbl.rows[r2].cells[num_col].text
    fmt_num = _infer_number_format(cell1_num, cell2_num)
    tmpl_row = tbl.rows[r2]

    insert_after = r2
    for k in range(3, min(MAX_ITEMS_FOR_TEMPLATE, len(items)) + 1):
        new_row = _insert_template_row_after(tbl, tmpl_row, insert_after)
        insert_after += 1

        try:
            new_row.cells[num_col].text = fmt_num(k)
            copy_cell_alignment(tmpl_row.cells[num_col], new_row.cells[num_col])
            enforce_times12_cell(new_row.cells[num_col])
        except Exception:
            pass

        ch2, chk = tag_n(TAG_CH, 2), tag_n(TAG_CH, k)
        pr2, prk = tag_n(TAG_PR, 2), tag_n(TAG_PR, k)
        am2, amk = tag_n(TAG_AM, 2), tag_n(TAG_AM, k)
        sd2, sdk = tag_n(TAG_SD, 2), tag_n(TAG_SD, k)

        if col_map['channel'] is not None:
            replace_tag_in_cell(new_row.cells[col_map['channel']], ch2, chk, tmpl_row.cells[col_map['channel']])
        else:
            for ci, cell in enumerate(new_row.cells):
                if ch2 in (cell.text or ""):
                    replace_tag_in_cell(cell, ch2, chk, tmpl_row.cells[ci])
                    break

        if col_map['period'] is not None:
            replace_tag_in_cell(new_row.cells[col_map['period']], pr2, prk, tmpl_row.cells[col_map['period']])
        else:
            for ci, cell in enumerate(new_row.cells):
                if pr2 in (cell.text or "") or "Срок размещения" in (cell.text or ""):
                    replace_tag_in_cell(cell, pr2, prk, tmpl_row.cells[ci if ci < len(tmpl_row.cells) else -1])
                    break

        if col_map['amount'] is not None:
            replace_tag_in_cell(new_row.cells[col_map['amount']], am2, amk, tmpl_row.cells[col_map['amount']])
        else:
            for ci, cell in enumerate(new_row.cells):
                if am2 in (cell.text or ""):
                    replace_tag_in_cell(cell, am2, amk, tmpl_row.cells[ci])
                    break

        if col_map['sdate'] is not None:
            replace_tag_in_cell(new_row.cells[col_map['sdate']], sd2, sdk, tmpl_row.cells[col_map['sdate']])
        else:
            for ci, cell in enumerate(new_row.cells):
                if sd2 in (cell.text or ""):
                    replace_tag_in_cell(cell, sd2, sdk, tmpl_row.cells[ci])
                    break

        for ci, cell in enumerate(new_row.cells):
            enforce_times12_cell(cell)
            if ci < len(tmpl_row.cells):
                copy_cell_alignment(tmpl_row.cells[ci], cell)


def render_docx_with_dynamic_rows(template_path: str, output_path: str, replacements: dict, items: list | None,
                                  enable_dynamic: bool) -> bool:
    try:
        if not os.path.exists(template_path):
            logging.error(f"❌ Шаблон не найден: {template_path}")
            return False
        doc = Document(template_path)
        if enable_dynamic and items and len(items) > 2:
            add_dynamic_rows_for_items(doc, items)

        body_map = dict(replacements)
        header_map = dict(replacements)
        if "{{CONTRACT_NUMBER}}" in replacements:
            header_map["<<CN>>"] = replacements["{{CONTRACT_NUMBER}}"]
        if "{{CONTRACT_DATE}}" in replacements:
            header_map["<<CD>>"] = replacements["{{CONTRACT_DATE}}"]
        if "{{INVOICE_NUMBER}}" in replacements:
            header_map["<<IN>>"] = replacements["{{INVOICE_NUMBER}}"]
        if "{{DATE}}" in replacements:
            header_map["<<DT>>"] = replacements["{{DATE}}"]

        _replace_in_block(doc, body_map)
        for section in doc.sections:
            _replace_in_header_footer(section.header, header_map)
            _replace_in_header_footer(section.footer, header_map)

        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    enforce_times12_cell(cell)

        doc.save(output_path)
        return True
    except Exception:
        logging.error("❌ Ошибка подстановки/динамики в DOCX")
        logging.error(traceback.format_exc())
        return False

# ================== ТЕКСТЫ ====================

INVOICE_PROMPTS = {
    "customer_name": (
        "1/5. Введите:\n"
        "***Наименование организации:***\n"
        "_(ООО/ИП/ФИО и т.п.)_\n\n"
        "_Пример: ООО «Показательный», ИП Круг Иван Иванович_"
    ),
    "customer_inn": (
        "2/5. Введите:\n"
        "***ИНН заказчика:***\n"
        "_(12 цифр для физ.лиц | 10 цифр для юр.лиц)_\n\n"
        "_Пример: 1236549876_"
    ),
    "item_channel": (
        "3/5. Укажите:\n"
        "***Где фактически размещается РИМ***\n"
        "(_Группа|Канал|Площадка|Ресурс этого пункта_)\n\n"
        "_Пример: Московская барыня_"
    ),
    "item_period": (
        "4/5. Укажите:\n"
        "***Срок оказания услуг | Период размещения РИМ*** для *этого пункта*:\n\n"
        "_Пример: с 03.11.2025 12:00 по 04.11.2025 12:00_"
    ),
    "item_amount": (
        "5/5. Укажите:\n"
        "***Стоимость услуг (руб.)*** для *этого пункта*:\n\n"
        "_Пример: 15000 или 15 000 или 15.000_\n"
        "❗*БЕЗ приписок ₽ / Р / руб. и т.п.*❗"
    ),
}


CONTRACT_PROMPTS = {
    ContractForm.customer_name: (
        "1/7. Введите:\n"
        "***Наименование организации заказчика:***\n"
        "_(ООО/ИП/ФИО и т.п.)_\n\n"
        "_Пример: ООО «РЕКЛАМА», ИП Иванов Иван Иванович_"
    ),
    ContractForm.customer_inn: (
        "2/7. Введите:\n"
        "***ИНН заказчика:***\n\n"
        "_Пример: 1236549876_\n"
        "_(12 цифр для физ.лиц | 10 цифр для юр.лиц)_"
    ),
    ContractForm.customer_ogrn: (
        "3/7. Введите:\n"
        "***ОГРН/ОГРНИП заказчика:***\n\n"
        "_Если нет — напишите «Нет/Отсутствует»_"
    ),
    ContractForm.placement_channel: (
        "4/7. Укажите:\n"
        "***Где фактически размещается РИМ***\n"
        "(_Группа|Канал|Площадка|Ресурс этого пункта_)\n\n"
        "_Пример: Питерский Гусь_"
    ),
    ContractForm.service_date: (
        "5/7. Укажите:\n"
        "***Дата размещения РИМ***\n\n"
        "_Пример: 03.11.2025 12:00_"
    ),
    ContractForm.service_period: (
        "6/7. Укажите:\n"
        "***Период|Срок размещения РИМ*** для *этого пункта*.\n\n"
        "_Пример: 24 часа с 03.11.2025 09:00 по 04.11.2025 09:00_"
    ),
    ContractForm.amount: (
        "7/7. Укажите:\n"
        "***Стоимость услуг (руб.)*** для *этого пункта*:\n\n"
        "_Пример: 15000 или 15 000 или 15.000_\n"
        "❗*БЕЗ приписок ₽ / Р / руб. и т.п.*❗"
    ),
}


def build_unified_caption(
    doc_kind: str,
    number: str,
    date_str: str,
    customer_name: str,
    inn: str,
    ogrn: str,
    service_start_date: str,
    period_text: str,
    services_count: int,
    total_sum_digits: int,
) -> str:
    """
    Единое резюме для всех типов документов (счёт-оферта и договор РИМ).

    doc_kind: "invoice" или "contract" – влияет только на первую строку.
    """
    num_e = md_escape(number)
    date_e = md_escape(date_str)
    name_e = md_escape(customer_name or "—")
    inn_e = md_escape(inn or "—")
    ogrn_e = md_escape(ogrn or "—")
    start_e = md_escape(service_start_date or "—")
    period_e = md_escape(period_text or "—")
    count = services_count or 0

    total = total_sum_digits or 0
    total_fmt = fmt_amount(total)
    words = number_to_words_ru(total)
    if words:
        words = words[:1].upper() + words[1:]
    words_e = md_escape(words)

    if doc_kind == "contract":
        first_line = f"🧾 Договор №РИМ/{num_e} от {date_e}"
    else:
        first_line = f"🧾 Счёт-оферта №{num_e} от {date_e}"

    caption = (
        f"*{first_line}*\n"
        f"Заказчик: {name_e}\n"
        f"ИНН: `{inn_e}`\n"
        f"ОГРН|ОГРНИП: {ogrn_e}\n"
        f"Период: {period_e}\n"
        f"╰⪼Кол-во услуг в ЭДО: {count} шт.\n\n"
        f"💲 *Общая сумма:* {total_fmt} ₽\n"
        f"_{words_e} руб., 00 коп._"
    )
    return caption[:CAPTION_LIMIT]

# ================== «Назад» ====================
async def handle_back(message: Message, state: FSMContext):
    """
    Кнопка «Назад» теперь НЕ теряет введённые ранее данные.
    Логика для счёта:
    - Если на этапе подтверждения (confirm) — извлекаем последний добавленный пункт из items,
      переносим его поля в временные _item_channel/_item_period и даём изменить сначала сумму.
      (То есть фактически редактируем последний пункт, а не добавляем новый.)
    - На шагах item_amount/item_period/item_channel — просто двигаемся назад без очистки уже введённых данных.
    """
    cur = await state.get_state()
    if not cur:
        await message.answer("Сейчас не в сценарии ввода. Выберите действие:", reply_markup=main_kb())
        return

    # —— СЧЁТ ——
    if cur == InvoiceForm.customer_inn.state:
        await state.set_state(InvoiceForm.customer_name)
        await message.answer(INVOICE_PROMPTS["customer_name"], reply_markup=step_kb())
        return

    if cur == InvoiceForm.item_channel.state:
        # Назад к ИНН без потери уже введённых ранее пунктов/items
        await state.set_state(InvoiceForm.customer_inn)
        await message.answer(INVOICE_PROMPTS["customer_inn"], reply_markup=step_kb())
        return

    if cur == InvoiceForm.item_period.state:
        # Назад к "канал", сохраняем _item_channel
        await state.set_state(InvoiceForm.item_channel)
        await message.answer(INVOICE_PROMPTS["item_channel"], reply_markup=step_kb())
        return

    if cur == InvoiceForm.item_amount.state:
        # Назад к "период", сохраняем _item_channel/_item_period
        await state.set_state(InvoiceForm.item_period)
        await message.answer(INVOICE_PROMPTS["item_period"], reply_markup=step_kb())
        return

    if cur == InvoiceForm.manual_text.state:
        # Отмена ввода ручного пункта — возвращаемся к выбору действий по счёту
        await state.set_state(InvoiceForm.confirm)
        await message.answer("Вернулся к выбору действий по счёту.", reply_markup=invoice_actions_kb())
        return

    if cur == InvoiceForm.manual_amount.state:
        # Назад к тексту ручного пункта
        await state.set_state(InvoiceForm.manual_text)
        await message.answer(
            "Введите произвольный текст, необходимый внести в пункт.\nНапример: Налог 5%",
            reply_markup=step_kb(),
        )
        return

    if cur == InvoiceForm.confirm.state:
        # Пользователь хочет отредактировать последний добавленный пункт.
        data = await state.get_data()
        items = list(data.get("items", []))
        if items:
            last = items.pop()  # убираем последний пункт, чтобы не дублировать при повторном вводе
            await state.update_data(
                items=items,
                _item_channel=last.get("channel", ""),
                _item_period=last.get("period", ""),
            )
        # Возвращаемся на этап ввода цены (можно ещё раз нажать «Назад»,
        # чтобы поправить период/канал — они уже в _item_*)
        await state.set_state(InvoiceForm.item_amount)
        await message.answer(INVOICE_PROMPTS["item_amount"], reply_markup=step_kb())
        return

    # —— ДОГОВОР ——
    if cur == ContractForm.customer_inn.state:
        await state.set_state(ContractForm.customer_name)
        await message.answer(CONTRACT_PROMPTS[ContractForm.customer_name], reply_markup=step_kb())
        return

    if cur == ContractForm.customer_ogrn.state:
        await state.set_state(ContractForm.customer_inn)
        await message.answer(CONTRACT_PROMPTS[ContractForm.customer_inn], reply_markup=step_kb())
        return

    if cur == ContractForm.placement_channel.state:
        await state.set_state(ContractForm.customer_ogrn)
        await message.answer(CONTRACT_PROMPTS[ContractForm.customer_ogrn], reply_markup=step_kb())
        return

    if cur == ContractForm.service_date.state:
        await state.set_state(ContractForm.placement_channel)
        await message.answer(CONTRACT_PROMPTS[ContractForm.placement_channel], reply_markup=step_kb())
        return

    if cur == ContractForm.service_period.state:
        await state.set_state(ContractForm.service_date)
        await message.answer(CONTRACT_PROMPTS[ContractForm.service_date], reply_markup=step_kb())
        return

    if cur == ContractForm.amount.state:
        await state.set_state(ContractForm.service_period)
        await message.answer(CONTRACT_PROMPTS[ContractForm.service_period], reply_markup=step_kb())
        return

    if cur == ContractForm.confirm.state:
        await state.set_state(ContractForm.amount)
        await message.answer(CONTRACT_PROMPTS[ContractForm.amount], reply_markup=step_kb())
        return

    # Если по какой-то причине состояние не распознано — не стираем данные!
    await message.answer("Возврат в главное меню.", reply_markup=main_kb())
    await state.clear()

# ================== ОСНОВНЫЕ ХЭНДЛЕРЫ ====================
async def cmd_start(message: Message, state: FSMContext):
    await state.clear()
    
    # Отслеживаем уникального пользователя
    user_id = message.from_user.id if message.from_user else None
    if user_id:
        is_new = track_unique_user(user_id)
        if is_new:
            logging.info(f"Новый уникальный пользователь: {user_id}")
    
    await message.answer(
        "_Запускаюсь..3..2..1.._\n🤖 𝙋𝙍𝙊𝙈𝙊-𝙋𝙍𝙊℠ _Рад..Снова..Видеть тебя!\n\nКакие задачи на сегодня?_\n*ВЫБЕРЕТЕ ДЕЙСТВИЕ В МЕНЮ*⤵︎",
        reply_markup=main_kb()
    )


async def cmd_feedback(message: Message, state: FSMContext):
    await state.clear()
    await message.answer("【Есть вопросы или предложения?】\n          Связаться ⌯⌲ @rusin_off", parse_mode=None)


async def offer_vk_lk_subscription(message: Message, state: FSMContext):
    # Сбрасываем текущее состояние и передаём управление в VK.ОРД-интеграцию
    await state.clear()
    await connect_vk_ord_lk(message, state)


async def show_subscription_terms(message: Message, state: FSMContext):
    text = (
        "           *ЧТО ТЫ ПОЛУЧАЕШЬ:*\n"
        "• *Автоматизация* _работы с рекламой и ОРД_\n"
        "• *ERID* _в один клик без ручного ввода_\n"
        "• *Доступ* _к \"базе данных\" аккаунта_\n"
        "• *Автоведение канала* _(удаление, открепление поста РИМ)_\n"
        "• *МЕНЬШЕ* _ошибок — _*БОЛЬШЕ* _времени на стратегию!_\n"
        "━━━━━━━━━━━━━━━━━\n"
        "          *ПОДРОБНЕЕ О ПОДПИСКЕ*\n"
        "                     ⬇️ ⬇️ ⬇️\n\n"
        "🧠 *Полная автоматизация* работы с рекламой\n"
        "_Планирование, подготовка, оформление рекламных постов — бот помогает делать всё быстрее и без хаоса._\n\n"
        "🔗 *Подключение кабинета* VK.ОРД прямо в боте\n"
        "_Больше никаких десятков вкладок и переключений между сервисами. Всё, что нужно для ОРД и маркировки, у тебя в одном месте — прямо в диалоге с ботом._\n\n"
        "🆔 *Получение ERID* в один клик\n"
        "_Забудь про ручной ввод и риск ошибиться. Нажал кнопку — получил корректный ERID. Бот помогает не пропустить ни один обязательный параметр._\n\n"
        "⚙️ *Расширенные функции* только для подписчиков\n"
        "_Умные подсказки по оформлению рекламы, ускорение подготовки постов, сохранённые шаблоны и история действий._\n"
        "_Бот превращается в полноценный рабочий инструмент, который реально экономит время и нервы._\n\n"
        "💸 *Всего 499 ₽ в месяц*\n"
        "_Это дешевле одной ошибки в рекламной кампании или одного «слива» бюджета из-за неправильных данных. Подписка окупается буквально одним предотвращённым косяком._\n"
        "━━━━━━━━━━━━━━━━━\n"
        "          *ПОЧЕМУ ЭТО ВЫГОДНО?*\n\n"
        "⏱️ *Экономия времени каждый день*\n"
        "_То, что ты обычно делаешь руками и по памяти, бот делает за тебя по кнопке. Меньше рутины — больше времени на креатив и переговоры._\n\n"
        "📈 *Чем БОЛЬШE рекламы — тем ВЫГОДНЕЕ подписка!*\n"
        "_Если у тебя уже есть обороты и постоянные размещения, подписка превращает работу в предсказуемый, быстрый конвейер:_ "
        "_ты занимаешься стратегией и деньгами — бот разбирается с рекламными постами, ОРД и ERID._\n\n"
        "💳 _Стоимость подписки_ - *499 ₽ / МЕС.*"
    )

    await message.answer(text, reply_markup=vk_lk_subscribe_kb())


async def ask_reset_confirmation(message: Message, state: FSMContext):
    text = ("Вы точно хотите сбросить порядковую нумерацию документов за сегодняшний день?")
    await message.answer(text, reply_markup=reset_confirm_kb())


async def cancel_reset_sequence(message: Message, state: FSMContext):
    await message.answer("Окей, нумерацию оставляю как есть. Возвращаюсь без изменений 👇", reply_markup=main_kb())


async def reset_sequence_cmd(message: Message, state: FSMContext):
    now = now_tz()
    uid = message.from_user.id if message.from_user else 0
    prev = reset_user_daily_sequence(now, uid)
    await state.clear()
    await message.answer(
        f"🔄 Последовательность на сегодня сброшена.\nБыло: {prev:02d} → Следующий номер будет: …-01",
        reply_markup=main_kb()
    )

# ——— СЧЁТ ———
async def start_invoice_flow(message: Message, state: FSMContext):
    await state.clear()
    await state.update_data(used_add_item=False, items=[])
    await state.set_state(InvoiceForm.customer_name)
    await message.answer(INVOICE_PROMPTS["customer_name"], reply_markup=step_kb())


async def cb_new_invoice(callback: CallbackQuery, state: FSMContext):
    await callback.answer()
    await start_invoice_flow(callback.message, state)


async def invoice_customer_name(message: Message, state: FSMContext):
    await state.update_data(customer_name=message.text.strip())
    await state.set_state(InvoiceForm.customer_inn)
    await message.answer(INVOICE_PROMPTS["customer_inn"], reply_markup=step_kb())


async def invoice_customer_inn(message: Message, state: FSMContext):
    await state.update_data(customer_inn=message.text.strip())
    await state.set_state(InvoiceForm.item_channel)
    await message.answer(
        "Теперь добавим первый пункт в счёт.\n\n" + INVOICE_PROMPTS["item_channel"],
        reply_markup=step_kb()
    )


async def add_item_start(message: Message, state: FSMContext):
    await state.update_data(used_add_item=True)
    await state.set_state(InvoiceForm.item_channel)
    await message.answer(INVOICE_PROMPTS["item_channel"], reply_markup=step_kb())


async def item_channel(message: Message, state: FSMContext):
    await state.update_data(_item_channel=message.text.strip())
    await state.set_state(InvoiceForm.item_period)
    await message.answer(INVOICE_PROMPTS["item_period"], reply_markup=step_kb())


async def item_period(message: Message, state: FSMContext):
    await state.update_data(_item_period=message.text.strip())
    await state.set_state(InvoiceForm.item_amount)
    await message.answer(INVOICE_PROMPTS["item_amount"], reply_markup=step_kb())


async def item_amount(message: Message, state: FSMContext):
    data = await state.get_data()
    channel = (data.get("_item_channel") or "").strip()
    period = (data.get("_item_period") or "").strip()
    amount_raw = message.text.strip()

    items = data.get("items", [])
    items.append({"channel": channel, "period": period, "amount": amount_raw})
    await state.update_data(items=items, _item_channel=None, _item_period=None)

    await message.answer(
        f"✅ Пункт добавлен:\n• Канал: {channel}\n• Период: {period}\n• Цена: {amount_raw}",
        reply_markup=invoice_actions_kb()
    )
    await state.set_state(InvoiceForm.confirm)


async def manual_pnc_start(message: Message, state: FSMContext):
    """
    Старт ручного добавления произвольного пункта (PNC).
    """
    await state.update_data(use_manual_pro_template=True)
    await state.set_state(InvoiceForm.manual_text)
    await message.answer(
        "Введите произвольный текст, необходимый внести в пункт.\nНапример: Налог 5%",
        reply_markup=step_kb()
    )


async def manual_pnc_text(message: Message, state: FSMContext):
    """
    Приём произвольного текстового описания пункта.
    """
    await state.update_data(manual_pnc_text=message.text.strip())
    await state.set_state(InvoiceForm.manual_amount)
    await message.answer("Введите сумму произвольного пункта", reply_markup=step_kb())


async def manual_pnc_amount(message: Message, state: FSMContext):
    """
    Приём суммы для произвольного пункта и возврат к экрану действий по счёту.
    """
    await state.update_data(manual_pnc_amount=message.text.strip())
    data = await state.get_data()
    txt = (data.get("manual_pnc_text") or "").strip()
    amount = (data.get("manual_pnc_amount") or "").strip()
    await message.answer(
        f"✅ Ручной пункт добавлен:\n• Описание: {txt}\n• Сумма: {amount}",
        reply_markup=invoice_actions_kb()
    )
    await state.set_state(InvoiceForm.confirm)


def build_invoice_caption_wrap(invoice_number, invoice_date, org_name, inn, positions_count, total_sum_digits,
                               total_sum_words):
    return build_unified_caption(
        doc_kind="invoice",
        number=invoice_number,
        date_str=invoice_date,
        customer_name=org_name,
        inn=inn,
        ogrn="",
        service_start_date="",
        period_text="",
        services_count=positions_count,
        total_sum_digits=total_sum_digits,
    )


async def form_invoice(message: Message, state: FSMContext, bot: Bot):
    data = await state.get_data()
    items = data.get("items", [])
    manual_pnc_text = (data.get("manual_pnc_text") or "").strip()
    manual_pnc_amount_raw = (data.get("manual_pnc_amount") or "").strip()

    # Без хотя бы одного стандартного пункта счёт не формируем
    if not items:
        await message.answer("Пока нет ни одного пункта. Сначала добавьте первый пункт.", reply_markup=step_kb())
        await state.set_state(InvoiceForm.item_channel)
        return

    now = now_tz()
    user_id = message.from_user.id if message.from_user else 0
    invoice_number = generate_number(now, user_id)
    invoice_date = generate_date(now)

    use_multi = bool(data.get("used_add_item")) or bool(data.get("use_manual_pro_template"))
    use_pro_template = bool(data.get("use_manual_pro_template"))

    if use_pro_template:
        template_path = TEMPLATE_INVOICE_MULTI_PRO
    else:
        template_path = TEMPLATE_INVOICE_MULTI if use_multi else TEMPLATE_INVOICE_SINGLE

    if not os.path.exists(template_path):
        alt = None
        if use_pro_template:
            if os.path.exists(TEMPLATE_INVOICE_MULTI):
                alt = TEMPLATE_INVOICE_MULTI
            elif os.path.exists(TEMPLATE_INVOICE_SINGLE):
                alt = TEMPLATE_INVOICE_SINGLE
        else:
            candidate = TEMPLATE_INVOICE_SINGLE if use_multi else TEMPLATE_INVOICE_MULTI
            if os.path.exists(candidate):
                alt = candidate

        if alt:
            await message.answer(
                f"ℹ️ Не нашёл шаблон:\n{template_path}\nИспользую альтернативный:\n{alt}",
                parse_mode=None
            )
            template_path = alt
        else:
            await message.answer("❌ Не найден ни один шаблон счёта.")
            return

    # Сумма по стандартным пунктам
    total_sum = sum(int(re.sub(r"[^\d]", "", i.get("amount") or "0") or 0) for i in items)

    # Плюсуем ручной пункт (если есть)
    manual_pnc_amount_int = 0
    if manual_pnc_amount_raw:
        manual_pnc_amount_int = int(re.sub(r"[^\d]", "", manual_pnc_amount_raw) or 0)
        total_sum += manual_pnc_amount_int

    total_sum_words = number_to_words_ru(total_sum)
    if total_sum_words:
        total_sum_words = total_sum_words[:1].upper() + total_sum_words[1:]

    first_service_date = (normalize_date_for_service_date(items[0].get("period", "")) if items else None) or invoice_date

    repl = {
        "{{INVOICE_NUMBER}}": invoice_number,
        "{{DATE}}": invoice_date,
        "{{CUSTOMER_NAME}}": data.get("customer_name", ""),
        "{{CUSTOMER_INN}}": data.get("customer_inn", ""),
        "{{TOTAL_SUM}}": fmt_amount(total_sum),
        "{{TOTAL_SUM_WORDS}}": total_sum_words,
        "{{AMOUNT_WORDS}}": total_sum_words,
        "{{SERVICE_DATE}}": first_service_date,
        "<<IN>>": invoice_number,
        "<<DT>>": invoice_date,
    }

    # Ручной пункт: текст и сумма в отдельные метки шаблона
    if manual_pnc_text:
        repl["{{PNC}}"] = manual_pnc_text
    if manual_pnc_amount_raw:
        repl["{{AMOUNT_PNC}}"] = manual_pnc_amount_raw

    for idx, item in enumerate(items, start=1):
        repl[tag_n("{{PLACEMENT_CHANNEL}}", idx)] = item.get("channel", "")
        repl[tag_n("{{SERVICE_PERIOD}}", idx)] = item.get("period", "")
        repl[tag_n("{{AMOUNT}}", idx)] = item.get("amount", "")
        sd_i = normalize_date_for_service_date(item.get("period", "")) or first_service_date
        repl[tag_n("{{SERVICE_DATE}}", idx)] = sd_i

    safe_name = (
        data.get("customer_name", "")
        .replace('"', "")
        .replace("«", "")
        .replace("»", "")
        .replace("/", "_")
        .replace("\\", "_")
        .replace(" ", "_")
    )[:50]
    output_path = os.path.join(OUTPUT_DIR, f"Счет-оферта_{safe_name}_{invoice_number}.docx")

    await message.answer("⏳ Формирую счёт…")
    ok = render_docx_with_dynamic_rows(
        template_path,
        output_path,
        replacements=repl,
        items=items,
        enable_dynamic=use_multi,
    )
    if not ok or not os.path.exists(output_path):
        await message.answer("❌ Не удалось создать счёт. Проверь шаблон и теги.")
        return

    period_main = items[0].get("period", "") if items else ""
    caption = build_unified_caption(
        doc_kind="invoice",
        number=invoice_number,
        date_str=invoice_date,
        customer_name=data.get("customer_name", "—"),
        inn=data.get("customer_inn", "—"),
        ogrn=data.get("customer_ogrn", ""),
        service_start_date=first_service_date,
        period_text=period_main,
        services_count=len(items),
        total_sum_digits=total_sum,
    )

    await bot.send_document(
        chat_id=message.chat.id,
        document=FSInputFile(output_path, filename=os.path.basename(output_path)),
        caption=caption,
        reply_markup=inline_new_invoice(),
    )
    await state.clear()


async def form_invoice_entry(message: Message, state: FSMContext, bot: Bot):
    await form_invoice(message, state, bot)

# ——— ДОГОВОР РИМ ———
async def start_contract_flow(message: Message, state: FSMContext):
    await state.clear()
    await state.update_data(used_add_item=False, items=[])
    await state.set_state(ContractForm.customer_name)
    await message.answer(CONTRACT_PROMPTS[ContractForm.customer_name], reply_markup=step_kb())


async def cb_new_contract(callback: CallbackQuery, state: FSMContext):
    await callback.answer()
    await start_contract_flow(callback.message, state)


async def contract_customer_name(message: Message, state: FSMContext):
    await state.update_data(customer_name=message.text.strip())
    await state.set_state(ContractForm.customer_inn)
    await message.answer(CONTRACT_PROMPTS[ContractForm.customer_inn], reply_markup=step_kb())


async def contract_customer_inn(message: Message, state: FSMContext):
    await state.update_data(customer_inn=message.text.strip())
    await state.set_state(ContractForm.customer_ogrn)
    await message.answer(CONTRACT_PROMPTS[ContractForm.customer_ogrn], reply_markup=step_kb())


async def contract_customer_ogrn(message: Message, state: FSMContext):
    await state.update_data(customer_ogrn=message.text.strip())
    await state.set_state(ContractForm.placement_channel)
    await message.answer(CONTRACT_PROMPTS[ContractForm.placement_channel], reply_markup=step_kb())


async def contract_placement_channel(message: Message, state: FSMContext):
    await state.update_data(_placement_channel=message.text.strip())
    await state.set_state(ContractForm.service_date)
    await message.answer(CONTRACT_PROMPTS[ContractForm.service_date], reply_markup=step_kb())


async def contract_service_date(message: Message, state: FSMContext):
    await state.update_data(_service_date=message.text.strip())
    await state.set_state(ContractForm.service_period)
    await message.answer(CONTRACT_PROMPTS[ContractForm.service_period], reply_markup=step_kb())


async def contract_service_period(message: Message, state: FSMContext):
    await state.update_data(_service_period=message.text.strip())
    await state.set_state(ContractForm.amount)
    await message.answer(CONTRACT_PROMPTS[ContractForm.amount], reply_markup=step_kb())


async def contract_amount(message: Message, state: FSMContext, bot: Bot):
    await state.update_data(_amount=message.text.strip())
    data = await state.get_data()
    item = {
        "channel": (data.get("_placement_channel") or "").strip(),
        "sdate": (data.get("_service_date") or "").strip(),
        "period": (data.get("_service_period") or "").strip(),
        "amount": (data.get("_amount") or "").strip(),
    }
    items = list(data.get("items", []))
    items.append(item)
    await state.update_data(
        items=items,
        _placement_channel=None,
        _service_date=None,
        _service_period=None,
        _amount=None
    )
    await message.answer(
        f"✅ Пункт добавлен:\n• Канал: {item['channel']}\n• Дата: {item['sdate']}\n• Период: {item['period']}\n• Цена: {item['amount']}",
        reply_markup=contract_actions_kb()
    )
    await state.set_state(ContractForm.confirm)


async def contract_add_item_start(message: Message, state: FSMContext):
    await state.update_data(used_add_item=True)
    await state.set_state(ContractForm.placement_channel)
    await message.answer(CONTRACT_PROMPTS[ContractForm.placement_channel], reply_markup=step_kb())


async def form_contract(message: Message, state: FSMContext, bot: Bot):
    data = await state.get_data()
    items = list(data.get("items", []))

    if not items:
        tmp = {
            "channel": (data.get("_placement_channel") or data.get("placement_channel") or "").strip(),
            "sdate": (data.get("_service_date") or data.get("service_date") or "").strip(),
            "period": (data.get("_service_period") or data.get("service_period") or "").strip(),
            "amount": (data.get("_amount") or data.get("amount") or "").strip(),
        }
        if any(tmp.values()):
            items = [tmp]

    if not items:
        await message.answer("Пока нет ни одного пункта. Сначала добавьте первый пункт.", reply_markup=step_kb())
        await state.set_state(ContractForm.placement_channel)
        return

    now = now_tz()
    user_id = message.from_user.id if message.from_user else 0
    contract_number = generate_number(now, user_id)
    contract_date = generate_date(now)

    use_multi = bool(data.get("used_add_item")) or (len(items) >= 2)
    template_path = TEMPLATE_CONTRACT_MULTI if use_multi else TEMPLATE_CONTRACT
    if not os.path.exists(template_path):
        alt = TEMPLATE_CONTRACT if use_multi else TEMPLATE_CONTRACT_MULTI
        if os.path.exists(alt):
            await message.answer(
                f"ℹ️ Не нашёл шаблон:\n{template_path}\nИспользую альтернативный:\n{alt}",
                parse_mode=None
            )
            template_path = alt
        else:
            await message.answer("❌ Не найден ни один шаблон договора.")
            return

    total_sum = 0
    norm_items = []
    for it in items:
        amt = only_digits(it.get("amount", ""))
        total_sum += amt
        norm_items.append({
            "channel": it.get("channel", ""),
            "period": it.get("period", ""),
            "amount": f"{amt:,}".replace(",", " "),
            "sdate": it.get("sdate", ""),
        })

    total_words = number_to_words_ru(total_sum) or ""
    if total_words:
        total_words = total_words[:1].upper() + total_words[1:]
    first_service_date = (normalize_date_for_service_date(items[0].get("period", "")) if items else None) or contract_date

    repl = {
        "{{CONTRACT_NUMBER}}": contract_number,
        "{{CONTRACT_DATE}}": contract_date,
        "{{CUSTOMER_NAME}}": data.get("customer_name", ""),
        "{{CUSTOMER_INN}}": data.get("customer_inn", ""),
        "{{CUSTOMER_OGRN}}": data.get("customer_ogrn", ""),
        "{{TOTAL_SUM}}": fmt_amount(total_sum),
        "{{TOTAL_SUM_WORDS}}": total_words,
        "{{AMOUNT_WORDS}}": total_words,
        "{{SERVICE_DATE}}": first_service_date,
        "<<CN>>": contract_number,
        "<<CD>>": contract_date,
    }
    for idx, item in enumerate(norm_items, start=1):
        repl[tag_n("{{PLACEMENT_CHANNEL}}", idx)] = item.get("channel", "")
        repl[tag_n("{{SERVICE_PERIOD}}", idx)] = item.get("period", "")
        repl[tag_n("{{AMOUNT}}", idx)] = item.get("amount", "")
        sd_i = normalize_date_for_service_date(item.get("period", "")) or item.get("sdate") or first_service_date
        repl[tag_n("{{SERVICE_DATE}}", idx)] = sd_i

    safe_name = (
        data.get("customer_name", "")
        .replace('"', "")
        .replace("'", "")
        .replace("/", "_")
        .replace("\\", "_")
        .replace(" ", "_")
    )[:50]
    output_path = os.path.join(OUTPUT_DIR, f"Договор_РИМ_{safe_name}_{contract_number}.docx")

    await message.answer("⏳ Формирую договор…")
    ok = render_docx_with_dynamic_rows(
        template_path,
        output_path,
        replacements=repl,
        items=norm_items,
        enable_dynamic=use_multi,
    )
    if not ok or not os.path.exists(output_path):
        await message.answer("❌ Не удалось создать договор. Проверьте шаблон и метки.")
        return

    period_main = items[0].get("period", "") if items else ""
    caption = build_unified_caption(
        doc_kind="contract",
        number=contract_number,
        date_str=contract_date,
        customer_name=data.get('customer_name', '—'),
        inn=data.get('customer_inn', '—'),
        ogrn=data.get('customer_ogrn', '—'),
        service_start_date=first_service_date,
        period_text=period_main,
        services_count=len(items),
        total_sum_digits=total_sum,
    )

    await bot.send_document(
        chat_id=message.chat.id,
        document=FSInputFile(output_path, filename=os.path.basename(output_path)),
        caption=caption,
        reply_markup=inline_new_contract(),
    )
    await state.clear()

# ——— Навигация ———
async def handle_cancel(message: Message, state: FSMContext):
    await state.clear()
    await message.answer("Окей, всё отменил. Возвращаю в главное меню 👇", reply_markup=main_kb())


async def vk_lk_no(message: Message, state: FSMContext):
    await handle_cancel(message, state)


# ================== МЕТРИКИ И СТАТИСТИКА ====================
async def cmd_stats(message: Message, state: FSMContext):
    """Команда для отправки статистики уникальных пользователей в админ чат."""
    await state.clear()
    
    if not ADMIN_CHAT_ID:
        await message.answer("❌ ADMIN_CHAT_ID не настроен в config.py")
        return
    
    try:
        stats = get_unique_users_stats()
        now = now_tz()
        date_str = now.strftime("%d.%m.%Y %H:%M")
        
        stats_text = (
            f"📊 *Статистика уникальных пользователей*\n"
            f"━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            f"👥 *Всего:* {stats['total']}\n"
            f"📅 *Сегодня:* {stats['today']}\n"
            f"📆 *За неделю:* {stats['week']}\n"
            f"📆 *За месяц:* {stats['month']}\n"
            f"━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            f"🕐 *Обновлено:* {date_str}"
        )
        
        # Отправляем в админ чат
        bot_instance = Bot(token=BOT_TOKEN)
        try:
            await bot_instance.send_message(
                chat_id=ADMIN_CHAT_ID,
                text=stats_text,
                parse_mode=ParseMode.MARKDOWN
            )
            await message.answer("✅ Статистика отправлена в админ чат")
        except Exception as e:
            await message.answer(f"❌ Ошибка отправки в админ чат: {str(e)}")
            logging.error(f"Ошибка отправки статистики в чат {ADMIN_CHAT_ID}: {e}")
        finally:
            await bot_instance.session.close()
            
    except Exception as e:
        await message.answer(f"❌ Ошибка получения статистики: {str(e)}")
        logging.error(f"Ошибка в cmd_stats: {e}")


async def send_stats_to_admin_chat():
    """Автоматическая отправка статистики в админ чат (можно вызывать по расписанию)."""
    if not ADMIN_CHAT_ID:
        logging.warning("ADMIN_CHAT_ID не настроен, статистика не будет отправлена")
        return
    
    try:
        stats = get_unique_users_stats()
        now = now_tz()
        date_str = now.strftime("%d.%m.%Y %H:%M")
        
        stats_text = (
            f"📊 *Статистика уникальных пользователей*\n"
            f"━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            f"👥 *Всего:* {stats['total']}\n"
            f"📅 *Сегодня:* {stats['today']}\n"
            f"📆 *За неделю:* {stats['week']}\n"
            f"📆 *За месяц:* {stats['month']}\n"
            f"━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            f"🕐 *Обновлено:* {date_str}"
        )
        
        bot_instance = Bot(token=BOT_TOKEN)
        try:
            await bot_instance.send_message(
                chat_id=ADMIN_CHAT_ID,
                text=stats_text,
                parse_mode=ParseMode.MARKDOWN
            )
            logging.info(f"Статистика отправлена в админ чат {ADMIN_CHAT_ID}")
        finally:
            await bot_instance.session.close()
            
    except Exception as e:
        logging.error(f"Ошибка отправки статистики в админ чат: {e}")


# ================== ЗАПУСК ====================
async def main() -> None:
    session = AiohttpSession(timeout=30)
    bot = Bot(
        token=BOT_TOKEN,
        default=DefaultBotProperties(parse_mode=ParseMode.MARKDOWN),
        session=session
    )
    dp = Dispatcher()

    # старт / меню
    dp.message.register(cmd_start, CommandStart())
    dp.message.register(cmd_feedback, match_contains("обратная связь"))
    dp.message.register(cmd_feedback, Command("feedback"))
    
    # статистика
    dp.message.register(cmd_stats, Command("stats"))
    dp.message.register(cmd_stats, match_contains("статистика"))

    dp.message.register(offer_vk_lk_subscription, match_contains("подключить кабинет"))
    dp.message.register(offer_vk_lk_subscription, match_contains("vk.орд"))
    dp.message.register(vk_lk_no, match_contains("не надо"))
    dp.message.register(show_subscription_terms, match_contains("подробнее"))

    # Назад / На главную
    dp.message.register(handle_back, match_contains("назад"))
    dp.message.register(handle_cancel, match_contains("главн"))

    # Сброс нумерации
    dp.message.register(ask_reset_confirmation, match_contains("сброс последовательности"))
    dp.message.register(ask_reset_confirmation, match_contains("сброс нумерац"))
    dp.message.register(reset_sequence_cmd, F.text == "✔️ ДА")
    dp.message.register(cancel_reset_sequence, F.text == "❌ НЕТ")

    # счёт
    dp.message.register(start_invoice_flow, match_contains("выставить сч"))
    dp.message.register(start_invoice_flow, match_contains("счёт на оплату"))
    dp.callback_query.register(cb_new_invoice, F.data == "new_invoice")
    dp.message.register(invoice_customer_name, InvoiceForm.customer_name)
    dp.message.register(invoice_customer_inn, InvoiceForm.customer_inn)
    dp.message.register(add_item_start, match_contains("добавить пункт"), InvoiceForm.confirm)
    dp.message.register(item_channel, InvoiceForm.item_channel)
    dp.message.register(item_period, InvoiceForm.item_period)
    dp.message.register(item_amount, InvoiceForm.item_amount)
    dp.message.register(manual_pnc_start, match_contains("добавить вручную"), InvoiceForm.confirm)
    dp.message.register(manual_pnc_text, InvoiceForm.manual_text)
    dp.message.register(manual_pnc_amount, InvoiceForm.manual_amount)
    dp.message.register(form_invoice_entry, match_contains("сформировать сч"))

    # договор РИМ
    dp.message.register(start_contract_flow, match_contains("договор рим"))
    dp.callback_query.register(cb_new_contract, F.data == "new_contract")
    dp.callback_query.register(inn_prev_page, F.data == "inn_prev")
    dp.callback_query.register(inn_next_page, F.data == "inn_next")
    dp.callback_query.register(inn_back_to_main, F.data == "inn_main")

    dp.message.register(contract_customer_name, ContractForm.customer_name)
    dp.message.register(contract_customer_inn, ContractForm.customer_inn)
    dp.message.register(contract_customer_ogrn, ContractForm.customer_ogrn)
    dp.message.register(contract_placement_channel, ContractForm.placement_channel)
    dp.message.register(contract_service_date, ContractForm.service_date)
    dp.message.register(contract_service_period, ContractForm.service_period)
    dp.message.register(contract_amount, ContractForm.amount)
    dp.message.register(contract_add_item_start, match_contains("добавить пункт"), ContractForm.confirm)
    dp.message.register(form_contract, match_contains("сформировать дог"))

    try:
        me = await bot.get_me()
        logging.info("Запускаю бота @%s (id=%s)", getattr(me, "username", "?"), getattr(me, "id", "?"))
        await bot.delete_webhook(drop_pending_updates=True)

        # ================== ПОИСК ПО ИНН ====================
        dp.message.register(start_inn_search, F.text.lower().contains("поиск по инн"))
        dp.message.register(handle_inn_input, StateFilter("awaiting_inn_search"), F.text)
        # =====================================================

        # ================== VK.ОРД ====================
        dp.message.register(connect_vk_ord_lk, F.text.lower().contains('перейти в кабинет'))
        dp.message.register(connect_vk_ord_lk, match_contains("подключить кабинет"))
        dp.message.register(connect_vk_ord_lk, F.text.lower().contains('сгенерировать erid'))
        dp.message.register(connect_vk_ord_lk, F.text.lower().contains('генерация erid'))

        dp.message.register(
            vk_ord_start_choice,
            F.text.in_([
                '✔️ Да',
                '❌ Нет',
                '❌ Не надо',
                '📚 Подробнее',
                '🔙 Назад',
                '🔙 В главное меню',
            ])
        )

        dp.message.register(save_vk_ord_token, StateFilter("vk_ord_token"), F.text)

               # верхний уровень VK.ОРД

        # ➕ Добавить контрагента (новое название кнопки)
        dp.message.register(vk_ord_add_contractor, match_contains("добавить контрагента"))
        dp.message.register(vk_ord_add_contractor, F.text.lower().contains('добавить контрагента'))

        # Поддержка старого текста, если где-то ещё остался
        dp.message.register(vk_ord_add_contractor, match_contains("внести контрагента"))
        dp.message.register(vk_ord_add_contractor, F.text.lower().contains('внести контрагента'))

        # 🖥️ Отправить договор в ЕРИР (новое название кнопки)
        dp.message.register(vk_ord_add_contract, match_contains("отправить договор"))
        dp.message.register(vk_ord_add_contract, F.text.lower().contains('отправить договор'))

        # Поддержка старого текста «Добавить договор» на всякий случай
        dp.message.register(vk_ord_add_contract, match_contains("добавить договор"))
        dp.message.register(vk_ord_add_contract, F.text.lower().contains('добавить договор'))

        # Креативы (как было)
        dp.message.register(vk_ord_add_creative, match_contains("креатив"))
        dp.message.register(vk_ord_add_creative, F.text.lower().contains('креатив'))



        # шаги мастера VK.ОРД — контрагент
        dp.message.register(vk_ord_person_type_step,   StateFilter("vk_ord_person_type"))
        dp.message.register(vk_ord_person_name_step,   StateFilter("vk_ord_person_name"))
        dp.message.register(vk_ord_person_inn_step,    StateFilter("vk_ord_person_inn"))
        dp.message.register(vk_ord_person_ogrn_step,   StateFilter("vk_ord_person_ogrn"))
        dp.message.register(vk_ord_person_roles_step,  StateFilter("vk_ord_person_roles"))
        dp.message.register(vk_ord_person_confirm_step, StateFilter("vk_ord_person_confirm"))

        # шаги мастера VK.ОРД — договор
        dp.message.register(vk_ord_additional_client_step,      StateFilter("vk_ord_additional_client"))
        dp.message.register(vk_ord_additional_contractor_step,  StateFilter("vk_ord_additional_contractor"))
        dp.message.register(vk_ord_additional_subject_step,     StateFilter("vk_ord_additional_subject"))
        dp.message.register(vk_ord_additional_date_step,        StateFilter("vk_ord_additional_date"))
        dp.message.register(vk_ord_additional_confirm_step,     StateFilter("vk_ord_additional_confirm"))

        dp.message.register(vk_ord_contract_type_step,  StateFilter("vk_ord_contract_type"))
        dp.message.register(vk_ord_contract_number_step,  StateFilter("vk_ord_contract_number"))
        dp.message.register(vk_ord_contract_date_step,    StateFilter("vk_ord_contract_date"))
        dp.message.register(vk_ord_contract_subject_step, StateFilter("vk_ord_contract_subject"))
        dp.message.register(vk_ord_contract_amount_step,  StateFilter("vk_ord_contract_amount"))
        dp.message.register(vk_ord_contract_confirm_step, StateFilter("vk_ord_contract_confirm"))

        dp.message.register(vk_ord_service_serial_step,      StateFilter("vk_ord_service_serial"))
        dp.message.register(vk_ord_service_comment_step,     StateFilter("vk_ord_service_comment"))
        dp.message.register(vk_ord_service_client_step,      StateFilter("vk_ord_service_client"))
        dp.message.register(vk_ord_service_contractor_step,  StateFilter("vk_ord_service_contractor"))
        dp.message.register(vk_ord_service_subject_step,     StateFilter("vk_ord_service_subject"))
        dp.message.register(vk_ord_service_date_step,        StateFilter("vk_ord_service_date"))
        dp.message.register(vk_ord_service_amount_step,      StateFilter("vk_ord_service_amount"))
        dp.message.register(vk_ord_service_confirm_step,     StateFilter("vk_ord_service_confirm"))

        # шаги мастера VK.ОРД — креатив
        dp.message.register(vk_ord_creative_name_step,    StateFilter("vk_ord_creative_name"))
        dp.message.register(vk_ord_creative_url_step,     StateFilter("vk_ord_creative_url"))
        dp.message.register(vk_ord_creative_period_step,  StateFilter("vk_ord_creative_period"))
        dp.message.register(vk_ord_creative_texts_step,   StateFilter("vk_ord_creative_texts"))
        dp.message.register(vk_ord_creative_media_step,   StateFilter("vk_ord_creative_media"))
        dp.message.register(vk_ord_creative_kktu_step,    StateFilter("vk_ord_creative_kktu"))
        dp.message.register(vk_ord_creative_confirm_step, StateFilter("vk_ord_creative_confirm"))
        # ==============================================



        await dp.start_polling(bot, allowed_updates=["message", "callback_query"])
    finally:
        await bot.session.close()

# ================== VK.ОРД ИНТЕГРАЦИЯ ====================
import json as _json_vk
import os as _os_vk
import io as _io_vk
import aiohttp as _aiohttp_vk
import time as _time_vk
import asyncio as _asyncio_vk
import re as _re_vk
from logging import getLogger as _getLogger_vk
from aiogram.types import Message as _Message_vk, ReplyKeyboardMarkup as _ReplyKeyboardMarkup_vk, KeyboardButton as _KeyboardButton_vk
from aiogram.fsm.context import FSMContext as _FSMContext_vk

VK_ORD_TOKENS_FILE = "secrets/vk_ord_tokens.json"
VK_ORD_STATE_FILE = "secrets/vk_ord_state.json"

# Переопределение через переменные окружения (опционально, приоритет выше config)
# Это позволяет переопределять настройки без изменения config.py (например, в Docker)
VK_ORD_API_BASE = _os_vk.getenv("VK_ORD_API_BASE", VK_ORD_API_BASE)
VK_ORD_API_TOKEN = _os_vk.getenv("VK_ORD_API_TOKEN", VK_ORD_API_TOKEN) if VK_ORD_API_TOKEN else None

if not VK_ORD_API_TOKEN:
    logging.warning("VK_ORD_API_TOKEN не задан! Функционал VK.ОРД может быть ограничен.")

# ---------- ХРАНЕНИЕ ТОКЕНОВ И СОСТОЯНИЯ ----------

def load_vk_ord_tokens() -> dict:
    if not _os_vk.path.exists(VK_ORD_TOKENS_FILE):
        return {}
    try:
        with open(VK_ORD_TOKENS_FILE, "r", encoding="utf-8") as f:
            return _json_vk.load(f)
    except Exception:
        return {}


def save_vk_ord_tokens(data: dict) -> None:
    with open(VK_ORD_TOKENS_FILE, "w", encoding="utf-8") as f:
        _json_vk.dump(data, f, ensure_ascii=False, indent=2)


def user_is_authorized(user_id: int | str) -> bool:
    tokens = load_vk_ord_tokens()
    return str(user_id) in tokens


def load_vk_ord_state() -> dict:
    if not _os_vk.path.exists(VK_ORD_STATE_FILE):
        return {}
    try:
        with open(VK_ORD_STATE_FILE, "r", encoding="utf-8") as f:
            return _json_vk.load(f)
    except Exception:
        return {}


def save_vk_ord_state(data: dict) -> None:
    with open(VK_ORD_STATE_FILE, "w", encoding="utf-8") as f:
        _json_vk.dump(data, f, ensure_ascii=False, indent=2)


def _get_user_state(user_id: str) -> dict:
    state = load_vk_ord_state()
    return state.get(user_id, {})


def _set_user_state(user_id: str, new_state: dict) -> None:
    state = load_vk_ord_state()
    state[user_id] = new_state
    save_vk_ord_state(state)


def _get_last_person(user_id: str) -> dict | None:
    st = _get_user_state(user_id)
    return st.get("last_person")


def _set_last_person(user_id: str, external_id: str, name: str, inn: str) -> None:
    st = _get_user_state(user_id)
    st["last_person"] = {"external_id": external_id, "name": name, "inn": inn}
    _set_user_state(user_id, st)

def _add_person_to_registry(user_id: str, external_id: str, name: str, inn: str) -> None:
    """
    Добавляем контрагента в локальный справочник бота для последующего поиска по названию или ИНН.
    """
    st = _get_user_state(user_id)
    persons = st.get("persons_registry", [])
    persons.append(
        {
            "external_id": external_id,
            "name": name,
            "inn": _re_vk.sub(r"\D", "", inn or ""),
        }
    )
    st["persons_registry"] = persons
    _set_user_state(user_id, st)


def _find_person_external_id(user_id: str, query: str) -> tuple[str | None, dict | None]:
    """
    Ищем external_id по названию или ИНН из локального справочника.
    Возвращаем (external_id, запись_контрагента) или (None, None).

    Приоритет поиска:
    1) Точное совпадение ИНН.
    2) Точное совпадение по названию (нормализованному).
    3) "Мягкий" поиск: по вхождению названия (нормализованного).
    """
    st = _get_user_state(user_id)
    persons = st.get("persons_registry", [])
    q = (query or "").strip().lower()
    inn_digits = _re_vk.sub(r"\D", "", q)

    def _norm_name(s: str) -> str:
        s = (s or "").lower()
        # убираем кавычки и лишнюю пунктуацию вокруг названия
        s = _re_vk.sub(r"[«»\"'“”„]", "", s)
        # схлопываем пробелы
        s = _re_vk.sub(r"\s+", " ", s).strip()
        return s

    q_norm = _norm_name(q)

    # 1. Сначала пробуем найти по ИНН — самое надёжное.
    if inn_digits:
        for p in persons:
            p_inn = _re_vk.sub(r"\D", "", p.get("inn") or "")
            if p_inn and p_inn == inn_digits:
                return p.get("external_id"), p

    # 2. Затем — точное совпадение названия (нормализованное).
    if q_norm:
        for p in persons:
            name_raw = (p.get("name") or "")
            name_norm = _norm_name(name_raw)
            if name_norm and name_norm == q_norm:
                return p.get("external_id"), p

    # 3. Мягкий поиск: вхождение строки.
    if q_norm:
        for p in persons:
            name_raw = (p.get("name") or "")
            name_norm = _norm_name(name_raw)
            if not name_norm:
                continue
            if q_norm in name_norm or name_norm in q_norm:
                return p.get("external_id"), p

    return None, None



def _get_last_contract(user_id: str) -> dict | None:
    st = _get_user_state(user_id)
    return st.get("last_contract")


def _set_last_contract(user_id: str, external_id: str, number: str, date: str) -> None:
    st = _get_user_state(user_id)
    st["last_contract"] = {"external_id": external_id, "number": number, "date": date}
    _set_user_state(user_id, st)


# ---------- КЛАВИАТУРЫ VK.ОРД ----------

def vk_lk_subscribe_kb() -> _ReplyKeyboardMarkup_vk:
    return _ReplyKeyboardMarkup_vk(
        keyboard=[
            [_KeyboardButton_vk(text="✔️ Да"), _KeyboardButton_vk(text="❌ Не надо")],
            [_KeyboardButton_vk(text="📚 Подробнее"), _KeyboardButton_vk(text="🔙 Назад")],
        ],
        resize_keyboard=True,
    )


def vk_ord_menu_kb() -> _ReplyKeyboardMarkup_vk:
    return _ReplyKeyboardMarkup_vk(
        keyboard=[
            [
                _KeyboardButton_vk(text="➕ Добавить контрагента"),
                _KeyboardButton_vk(text="🖥️ Отправить договор в ЕРИР"),
            ],
            [_KeyboardButton_vk(text="🖌️ Оформить креатив (ERID)")],
            [_KeyboardButton_vk(text="🔙 В главное меню")],
        ],
        resize_keyboard=True,
    )


def step_kb() -> _ReplyKeyboardMarkup_vk:
    return _ReplyKeyboardMarkup_vk(
        keyboard=[
            [_KeyboardButton_vk(text="◀  Назад"), _KeyboardButton_vk(text="✖  На главную")],
        ],
        resize_keyboard=True,
    )


# ---------- ПОДКЛЮЧЕНИЕ ЛК VK.ОРД ----------

async def connect_vk_ord_lk(message: _Message_vk, state: _FSMContext_vk):
    user_id = str(message.from_user.id)
    if user_is_authorized(user_id):
        await message.answer(
            "Ой, кажется Вы уже авторизовались..\n\n"
            "Выберите действие, которое хотите совершить в «VK.ОРД»:",
            reply_markup=vk_ord_menu_kb(),
        )
        return

    text = (
        "Хотите подключить личный кабинет в «VK.ОРД»?\n\n"
        "Доступ к сервису станет доступным ТОЛЬКО В ПЛАТНОЙ ВЕРСИИ бота\n\n"
        "Желаете перейти на платную версию и открыть все возможности?"
    )
    await message.answer(text, reply_markup=vk_lk_subscribe_kb())


async def vk_ord_start_choice(message: _Message_vk, state: _FSMContext_vk):
    txt = (message.text or "").strip()
    if txt == "✔️ Да":
        await message.answer(
            "Пожалуйста, отправьте ваш API-токен VK.ОРД.\n\n"
            "🔐 Найти его можно в личном кабинете VK.ОРД в разделе настроек/интеграций "
            "(см. официальную документацию)."
        )
        await state.set_state("vk_ord_token")
    elif txt == "📚 Подробнее":
        await message.answer(
            "ℹ️ Раздел VK.ОРД позволяет автоматически передавать сведения о рекламе, "
            "контрагентах, договорах и креативах в соответствии с законодательством РФ."
        )
    elif txt in {"❌ Нет", "❌ Не надо", "🔙 Назад", "🔙 В главное меню"}:
        await state.clear()
        await message.answer("Возвращаю вас в главное меню 😊")


async def save_vk_ord_token(message: _Message_vk, state: _FSMContext_vk):
    token = (message.text or "").strip()
    user_id = str(message.from_user.id)
    tokens = load_vk_ord_tokens()
    tokens[user_id] = token
    save_vk_ord_tokens(tokens)
    await state.clear()
    await message.answer("🎉 Поздравляю! Теперь Вы подключены к VK.ОРД.")
    await message.answer("Выберите действие, которое хотите совершить:", reply_markup=vk_ord_menu_kb())


# ---------- ОБЩИЙ КЛИЕНТ VK.ОРД API ----------


async def vk_ord_api_request(user_id: str, method: str, path: str | list, json_body: dict | None = None):
    """
    Универсальный помощник для вызова VK.ОРД API.

    ВНИМАНИЕ:
    1) Схему авторизации (обычно `Authorization: Bearer <TOKEN>`) нужно
       проверить по официальной документации VK.ОРД.
    2) Конкретные пути (`/v1/person/{external_id}`, `/v1/contract/{external_id}`,
       `/v3/creative/{external_id}` и т.п.) и структуру json_body нужно сверить
       со swagger-документацией VK.ОРД (sandbox/prod).
    """
    log = _getLogger_vk(__name__)
    tokens = load_vk_ord_tokens()
    # Пытаемся сначала взять персональный токен пользователя, если он сохранён,
    # иначе используем глобальный VK_ORD_API_TOKEN.
    token = tokens.get(str(user_id)) or VK_ORD_API_TOKEN
    if not token:
        return False, "API-токен VK.ОРД для этого пользователя не найден. Переподключите кабинет."

    base_raw = VK_ORD_API_BASE.rstrip("/")
    if not base_raw:
        return False, "Базовый URL VK.ОРД API не настроен. Установите VK_ORD_API_BASE."

    # Собираем относительный путь вида "v3/creative/{external_id}"
    if isinstance(path, str):
        rel_path = path.lstrip("/")
    else:
        rel_path = "/".join(str(p).strip("/") for p in path if p)

    url = base_raw.rstrip("/") + "/" + rel_path

    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json",
        "Accept": "application/json",
    }

    async def _do(session, url: str):
        async with session.request(method.upper(), url, json=json_body, headers=headers) as resp:
            txt = await resp.text()
            try:
                data = await resp.json()
            except Exception:
                data = None
            return resp.status, txt, data, url, dict(resp.headers)

    async with _aiohttp_vk.ClientSession() as session:
        last = None
        backoff = 0
        for attempt in range(3):
            if backoff:
                await _asyncio_vk.sleep(backoff)

            status, txt, data, used, resp_headers = await _do(session, url)

            if status == 429:
                ra = None
                if isinstance(resp_headers, dict):
                    ra = resp_headers.get("Retry-After") or resp_headers.get("retry-after")
                try:
                    backoff = max(1, int(ra)) if ra else (2 ** attempt)
                except Exception:
                    backoff = 2 ** attempt
                last = (status, txt, data, used)
                continue

            if 500 <= status < 600:
                backoff = 2 ** attempt
                last = (status, txt, data, used)
                continue

            if 200 <= status < 300:
                return True, data or txt

            last = (status, txt, data, used)
            break

        if last:
            status, txt, data, used = last
            log.error(
                "VK.ОРД API error: status=%s url=%s body=%r json=%r",
                status, used, txt, data
            )
            return False, data or txt or f"HTTP {status}"
        return False, "Не удалось вызвать VK.ОРД API: пустой ответ/нет попыток."
def _normalize_roles_to_codes(text: str) -> list[str]:
    """
    Преобразует человекочитаемые роли в коды ролей VK.ОРД.

    Поддерживаемые роли (см. примеры из документации VK.ОРД):
    - "рекламодатель"  -> "advertiser"
    - "площадка", "распространитель" -> "publisher"
    - "агентство", "агент", "посредник" -> "agency"
    - "ОРД" -> "ors"

    Пользователь может ввести роли через запятую, напр.:
    "Рекламодатель, Агентство".
    Неизвестные значения игнорируются.
    """
    if not text:
        return []

    parts = [p.strip().lower() for p in text.split(",") if p.strip()]
    result: list[str] = []

    for p in parts:
        code = None
        if "рекламод" in p:
            code = "advertiser"
        elif "распростран" in p or "площад" in p:
            code = "publisher"
        elif "агент" in p or "посред" in p:
            code = "agency"
        elif "орд" in p:
            code = "ors"

        if code and code not in result:
            result.append(code)

    return result


# ---------- МАСТЕР СОЗДАНИЯ КОНТРАГЕНТА ----------

def vk_ord_contractor_type_kb() -> _ReplyKeyboardMarkup_vk:
    """
    Клавиатура выбора типа контрагента для VK.ОРД.
    """
    return _ReplyKeyboardMarkup_vk(
        keyboard=[
            [
                _KeyboardButton_vk(text="Физ. лицо 👤"),
                _KeyboardButton_vk(text="Юр. лицо 🏢"),
                _KeyboardButton_vk(text="ИП            💼"),
            ],
            [
                _KeyboardButton_vk(text="◀  Назад"),
                _KeyboardButton_vk(text="✖  На главную"),
            ],
        ],
        resize_keyboard=True,
    )




def vk_ord_contract_type_kb() -> _ReplyKeyboardMarkup_vk:
    """
    Клавиатура выбора типа договора для VK.ОРД.
    Визуально повторяет меню выбора типа контрагента, но с вариантами договора.
    """
    return _ReplyKeyboardMarkup_vk(
        keyboard=[
            [
                _KeyboardButton_vk(text="🛠️           Оказание услуг           🛠️"),
            ],
            [
                _KeyboardButton_vk(text="Посредничество"),
                _KeyboardButton_vk(text="Доп. соглашения"),
            ],
            [
                _KeyboardButton_vk(text="◀  Назад"),
                _KeyboardButton_vk(text="✖  На главную"),
            ],
        ],
        resize_keyboard=True,
    )




async def _vk_ord_extract_telegram_media(message: _Message_vk):
    """
    Достаёт ОДИН медиафайл из сообщения Telegram и возвращает (bytes, filename, content_type),
    либо None, если медиа нет.
    """
    file_obj = None
    filename = "media.bin"
    content_type = "application/octet-stream"

    if message.document:
        file_obj = message.document
        filename = message.document.file_name or "document"
        content_type = message.document.mime_type or "application/octet-stream"
    elif message.photo:
        # Берём самое большое фото
        file_obj = message.photo[-1]
        filename = f"photo_{file_obj.file_unique_id}.jpg"
        content_type = "image/jpeg"
    elif message.video:
        file_obj = message.video
        filename = message.video.file_name or "video.mp4"
        content_type = message.video.mime_type or "video/mp4"
    elif message.animation:
        file_obj = message.animation
        filename = message.animation.file_name or "animation.gif"
        content_type = message.animation.mime_type or "image/gif"
    elif message.voice:
        file_obj = message.voice
        filename = "voice.ogg"
        content_type = "audio/ogg"
    elif message.audio:
        file_obj = message.audio
        filename = message.audio.file_name or "audio.mp3"
        content_type = message.audio.mime_type or "audio/mpeg"

    if not file_obj:
        return None

    buf = _io_vk.BytesIO()
    await message.bot.download(file_obj, buf)
    buf.seek(0)
    return buf.read(), filename, content_type


async def vk_ord_upload_media(
    user_id: str,
    file_bytes: bytes,
    filename: str,
    content_type: str,
):
    """
    Загружает медиафайл в VK.ОРД через PUT /v1/media/{external_id}.

    Возвращает (ok: bool, result),
    где result = external_id (если ok=True) или тело ошибки/ответа.
    """
    log = _getLogger_vk(__name__)
    tokens = load_vk_ord_tokens()
    token = tokens.get(str(user_id)) or VK_ORD_API_TOKEN

    base_raw = VK_ORD_API_BASE.rstrip("/")
    if not base_raw:
        return False, "Базовый URL VK.ОРД API не настроен. Установите VK_ORD_API_BASE."

    external_id = f"media-{int(_time_vk.time())}-{user_id}".replace(" ", "")
    url = f"{base_raw}/v1/media/{external_id}"

    form = _aiohttp_vk.FormData()
    form.add_field(
        "media_file",
        file_bytes,
        filename=filename,
        content_type=content_type or "application/octet-stream",
    )

    headers = {
        "Authorization": f"Bearer {token}",
    }

    async with _aiohttp_vk.ClientSession() as session:
        async with session.put(url, data=form, headers=headers) as resp:
            txt = await resp.text()
            try:
                data = await resp.json()
            except Exception:
                data = None

            if 200 <= resp.status < 300:
                if isinstance(data, dict):
                    eid = data.get("external_id") or data.get("id") or external_id
                else:
                    eid = external_id
                log.info("VK.ОРД media uploaded: status=%s url=%s external_id=%s", resp.status, url, eid)
                return True, eid

            log.error(
                "VK.ОРД media upload error: status=%s url=%s body=%r json=%r",
                resp.status, url, txt, data
            )
            return False, data or txt or f"HTTP {resp.status}"

async def vk_ord_add_contractor(message: _Message_vk, state: _FSMContext_vk):
    """
    Точка входа из меню бота в мастер создания контрагента VK.ОРД.
    На этом шаге пользователь выбирает тип контрагента.
    """
    user_id = str(message.from_user.id)
    if not user_is_authorized(user_id):
        await message.answer(
            "Сначала подключите личный кабинет VK.ОРД через главное меню.",
            reply_markup=vk_lk_subscribe_kb(),
        )
        return

    await state.clear()
    await state.set_state("vk_ord_person_type")
    await message.answer(
        "Запускаю ЕРИР..  _Анализируем рекламу.._\n"
        "Выберите *ТИП* Контрагента(Заказчика):",
        reply_markup=vk_ord_contractor_type_kb(),
    )



async def vk_ord_person_type_step(message: _Message_vk, state: _FSMContext_vk):
    """
    Обработка выбора типа контрагента.
    Для "Физ. лицо", "Юр. лицо" и "ИП" запускается общий мастер с разными подсказками
    и сохранением выбранного типа в состоянии.
    """
    # Нормализуем текст: убираем пробелы, эмодзи и другие не-буквенные символы в начале
    raw_text = (message.text or "").strip().lower()
    # отрезаем все не-«словесные» символы в начале (эмодзи, знаки и т.п.)
    text = _re_vk.sub(r"^[^\w]+", "", raw_text)

    if text.startswith("физ"):
        # Запуск мастера создания контрагента — физическое лицо
        await state.update_data(vk_ord_person_kind="physical")
        await state.set_state("vk_ord_person_name")
        await message.answer(
            "🧾 *Создание контрагента (ФИЗ. ЛИЦО, шаг 1/4)*\n\n"
            "Введите полное ФИО контрагента так, как оно указано в документах.\n"
            "Пример: *Иванов Сергей Петрович*.",
            reply_markup=step_kb(),
            parse_mode="Markdown",
        )
    elif text.startswith("юр"):
        # Запуск мастера создания контрагента — юридическое лицо
        await state.update_data(vk_ord_person_kind="juridical")
        await state.set_state("vk_ord_person_name")
        await message.answer(
            "🧾 *Создание контрагента (ЮР. ЛИЦО, шаг 1/4)*\n\n"
            "Введите полное наименование организации так, как оно указано в учредительных документах.\n"
            "Пример: *ООО «Вкусные десерты»*.",
            reply_markup=step_kb(),
            parse_mode="Markdown",
        )
    elif text.startswith("ип"):
        # Запуск мастера создания контрагента — индивидуальный предприниматель
        await state.update_data(vk_ord_person_kind="ip")
        await state.set_state("vk_ord_person_name")
        await message.answer(
            "🧾 *Создание контрагента (ИП, шаг 1/4)*\n\n"
            "Введите наименование контрагента в формате ИП, как в документах.\n"
            "Пример: *ИП Максимов Иван Семёнович*.",
            reply_markup=step_kb(),
            parse_mode="Markdown",
        )
    else:
        # Повторный запрос выбора типа
        await message.answer(
            "Пожалуйста, выберите один из вариантов: *👤 Физ. лицо*, *🏢 Юр. лицо* или *💼 ИП*.",
            parse_mode="Markdown",
            reply_markup=vk_ord_contractor_type_kb(),
        )

async def vk_ord_person_name_step(message: _Message_vk, state: _FSMContext_vk):
    name = (message.text or "").strip()
    if not name:
        await message.answer("Наименование не должно быть пустым, попробуйте ещё раз.", reply_markup=step_kb())
        return

    await state.update_data(vk_ord_person_name=name)
    await state.set_state("vk_ord_person_inn")
    await message.answer(
        "🧾 *Создание контрагента (шаг 2/4)*\n\n"
        "Укажите ИНН контрагента (10 или 12 цифр).",
        reply_markup=step_kb(),
        parse_mode="Markdown",
    )


async def vk_ord_person_inn_step(message: _Message_vk, state: _FSMContext_vk):
    inn = _re_vk.sub(r"\D", "", (message.text or ""))
    if not inn.isdigit() or len(inn) not in (10, 12):
        await message.answer(
            "ИНН должен содержать 10 или 12 цифр. Отправьте корректное значение.",
            reply_markup=step_kb(),
        )
        return

    await state.update_data(vk_ord_person_inn=inn)
    await state.set_state("vk_ord_person_ogrn")
    await message.answer(
        "🧾 *Создание контрагента (шаг 3/4)*\n\n"
        "Укажите ОГРН/ОГРНИП.\n"
        "Если нет или не хотите указывать — напишите «нет».",
        reply_markup=step_kb(),
        parse_mode="Markdown",
    )


async def vk_ord_person_ogrn_step(message: _Message_vk, state: _FSMContext_vk):
    ogrn_raw = (message.text or "").strip()
    ogrn = "" if ogrn_raw.lower() == "нет" else ogrn_raw
    await state.update_data(vk_ord_person_ogrn=ogrn)
    await state.set_state("vk_ord_person_roles")
    await message.answer(
        "🧾 *Создание контрагента (шаг 4/4)*\n\n"
        "Укажите роли контрагента в цепочке ОРД.\n"
        "Например: _Рекламодатель, Агентстсво, Рекламана система, Издатель._\n"
        "Можно несколько через запятую.",
        reply_markup=step_kb(),
        parse_mode="Markdown",
    )


async def vk_ord_person_roles_step(message: _Message_vk, state: _FSMContext_vk):
    roles_raw = (message.text or "").strip()
    await state.update_data(vk_ord_person_roles_raw=roles_raw)

    data = await state.get_data()
    text = (
        "Проверьте данные контрагента:\n"
        f"• Наименование: *{data.get('vk_ord_person_name', '')}*\n"
        f"• ИНН: `{data.get('vk_ord_person_inn', '')}`\n"
        f"• ОГРН/ОГРНИП: `{data.get('vk_ord_person_ogrn', '') or '—'}`\n"
        f"• Роли: {roles_raw or 'не указаны'}\n\n"
        "Если всё верно — нажмите «✅ Подтвердить».\n"
        "Чтобы отменить — используйте «✖  На главную»."
    )
    kb = _ReplyKeyboardMarkup_vk(
        keyboard=[
            [_KeyboardButton_vk(text="✅ Подтвердить")],
            [_KeyboardButton_vk(text="◀  Назад"), _KeyboardButton_vk(text="✖  На главную")],
        ],
        resize_keyboard=True,
    )
    await state.set_state("vk_ord_person_confirm")
    await message.answer(text, reply_markup=kb, parse_mode="Markdown")



async def vk_ord_person_confirm_step(message: _Message_vk, state: _FSMContext_vk):
    """
    Финальное подтверждение создания контрагента VK.ОРД.

    Работает для "Физ. лица", "Юр. лица" и "ИП":
    выбранный тип сохраняется в состоянии как vk_ord_person_kind и
    влияет на поле juridical_details.type в payload.
    """
    if (message.text or "").strip() != "✅ Подтвердить":
        await message.answer(
            "Чтобы создать контрагента, нажмите «✅ Подтвердить» или используйте «◀  Назад»/«✖  На главную».",
            reply_markup=step_kb(),
        )
        return

    data = await state.get_data()
    user_id = str(message.from_user.id)
    ext_id = f"tg-{user_id}-person-{int(_time_vk.time())}"

    roles_codes = _normalize_roles_to_codes(data.get("vk_ord_person_roles_raw", ""))
    # Если ни одной роли распознать не удалось, по умолчанию считаем контрагента рекламодателем.
    if not roles_codes:
        roles_codes = ["advertiser"]

    # Определяем тип контрагента
    kind = (data.get("vk_ord_person_kind") or "").strip().lower()
    inn_raw = (data.get("vk_ord_person_inn") or "").strip()
    inn_digits = _re_vk.sub(r"\D", "", inn_raw)

    # Используем настройки типов персон из config
    if kind == "juridical":
        _person_type = VK_ORD_PERSON_TYPE_JURIDICAL
    elif kind == "ip":
        # ИП — отдельный тип в VK.ОРД (см. пример person/type=ip)
        _person_type = VK_ORD_PERSON_TYPE_IP
    elif kind == "physical":
        # Физ. лицо — резервная логика по длине ИНН
        if len(inn_digits) == 10:
            _person_type = VK_ORD_PERSON_TYPE_JURIDICAL
        elif len(inn_digits) == 12:
            _person_type = VK_ORD_PERSON_TYPE_INDIVIDUAL
        else:
            _person_type = VK_ORD_PERSON_TYPE_DEFAULT
    else:
        # На всякий случай используем определение по длине ИНН
        if len(inn_digits) == 10:
            _person_type = VK_ORD_PERSON_TYPE_JURIDICAL
        elif len(inn_digits) == 12:
            _person_type = VK_ORD_PERSON_TYPE_INDIVIDUAL
        else:
            _person_type = VK_ORD_PERSON_TYPE_DEFAULT

    if not inn_digits:
        inn_digits = inn_raw

    payload = {
        "name": data.get("vk_ord_person_name"),
        "roles": roles_codes,
        "juridical_details": {
            "type": _person_type,
            "inn": inn_digits,
        },
    }

    ogrn_val = (data.get("vk_ord_person_ogrn") or "").strip()
    if ogrn_val and ogrn_val.lower() != "нет":
        payload["juridical_details"]["ogrn"] = ogrn_val

    ok, resp = await vk_ord_api_request(user_id, "PUT", f"/v1/person/{ext_id}", payload)
    if not ok:
        await message.answer(
            "❌ Не удалось создать контрагента через VK.ОРД API.\n\n"
            f"*Ответ сервера:* `{resp}`\n\n"
            "Сверьтесь с документацией VK.ОРД по методу создания контрагента и скорректируйте данные.",
            parse_mode="Markdown",
            reply_markup=vk_ord_menu_kb(),
        )
        await state.clear()
        return

    _set_last_person(user_id, ext_id, data.get("vk_ord_person_name", ""), inn_digits)
    _add_person_to_registry(user_id, ext_id, data.get("vk_ord_person_name", ""), inn_digits)
    text = "✅ Контрагент успешно создан и *отправлен в ЕРИР* на проверку!\n"
    if isinstance(resp, dict):
        vk_id = resp.get("id")
        if vk_id:
            text += f"ID в VK.ОРД: `{vk_id}`\n"
    text += f"_Пожалуйста, проверьте Ваш личный кабинет._ `{ext_id}`"
    await message.answer(text, parse_mode="Markdown", reply_markup=vk_ord_menu_kb())
    await state.clear()


async def vk_ord_add_contract(message: _Message_vk, state: _FSMContext_vk):
    """
    Новый вход в мастер добавления договора VK.ОРД.

    Вместо немедленного запуска старого мастера договоров показывает
    меню выбора типа договора (Оказание услуг / Посредничество / Дополнительное соглашение).
    Пока все типы работают как заглушки, но сама развилка остаётся для дальнейшей доработки.
    """
    user_id = str(message.from_user.id)
    if not user_is_authorized(user_id):
        await message.answer(
            "Сначала подключите личный кабинет VK.ОРД через главное меню.",
            reply_markup=vk_lk_subscribe_kb(),
        )
        return

    # При необходимости можно будет проверить наличие контрагента,
    # но сейчас просто показываем выбор типа договора.
    await state.clear()
    await state.set_state("vk_ord_contract_type")
    await message.answer(
        "Запускаю ЕРИР..  _Подгружаю ERID'ы.._\n"
        "Пожалуйста, выберите *ТИП* договора:",
        reply_markup=vk_ord_contract_type_kb(),
    )




async def vk_ord_contract_type_step(message: _Message_vk, state: _FSMContext_vk):
    """
    Обработка выбора типа договора VK.ОРД.

    Для "Оказание услуг" запускается мастер договора типа service.
    Для "Посредничество" пока остаётся заглушка.
    Для "Доп. соглашения" запускается мастер создания договора типа additional.
    """
    raw_text = (message.text or "").strip().lower()
    # Отрезаем возможные эмодзи/знаки в начале (на будущее)
    text = _re_vk.sub(r"^[^\w]+", "", raw_text)

    # Оказание услуг — полноценный мастер service
    if text.startswith("оказан"):
        await state.update_data(vk_ord_contract_kind="service")
        await state.set_state("vk_ord_service_serial")
        await message.answer(
            "🧾 *Создание договора (Оказание услуг, шаг 1/7)*\n\n"
            "Укажите серийный номер договора.",
            reply_markup=step_kb(),
            parse_mode="Markdown",
        )
        return

    # Посредничество — пока заглушка
    if text.startswith("посред"):
        human_readable = "Посредничество"
        await state.clear()
        await message.answer(
            (
                f"Увы, тип договора «{human_readable}» пока недоступен.\n"
                "Функционал в разработке — скоро здесь появится мастер создания договора по API VK.ОРД.\n\n"
                "Возвращаю вас в меню VK.ОРД."
            ),
            reply_markup=vk_ord_menu_kb(),
        )
        return

    # Доп. соглашения — полноценный мастер additional
    if text.startswith("доп"):
        await state.update_data(vk_ord_contract_kind="additional")
        await state.set_state("vk_ord_additional_client")
        await message.answer(
            "🧾 *Создание доп. соглашения (шаг 1/4)*\n\n"
            "Укажите `external_id` заказчика (`client_external_id`) для доп. соглашения.",
            reply_markup=step_kb(),
            parse_mode="Markdown",
        )
        return

    # Некорректный ввод — повторяем выбор
    await message.answer(
        "Пожалуйста, выберите ТИП договора, используя кнопки ниже.",
        reply_markup=vk_ord_contract_type_kb(),
    )


async def vk_ord_additional_client_step(message: _Message_vk, state: _FSMContext_vk):
    """
    Шаг 1/4: заказчик (client_external_id) для доп. соглашения.
    """
    client_ext_id = (message.text or "").strip()
    if not client_ext_id:
        await message.answer(
            "Поле `external_id` заказчика не должно быть пустым. Укажите значение или вернитесь назад.",
            reply_markup=step_kb(),
        )
        return

    await state.update_data(vk_ord_additional_client_external_id=client_ext_id)
    await state.set_state("vk_ord_additional_contractor")
    await message.answer(
        "🧾 *Создание доп. соглашения (шаг 2/4)*\n\n"
        "Укажите `external_id` исполнителя (он же издатель) — `contractor_external_id`.",
        reply_markup=step_kb(),
        parse_mode="Markdown",
    )


async def vk_ord_additional_contractor_step(message: _Message_vk, state: _FSMContext_vk):
    """
    Шаг 2/4: исполнитель (contractor_external_id) для доп. соглашения.
    """
    contractor_ext_id = (message.text or "").strip()
    if not contractor_ext_id:
        await message.answer(
            "Поле `external_id` исполнителя не должно быть пустым. Укажите значение или вернитесь назад.",
            reply_markup=step_kb(),
        )
        return

    await state.update_data(vk_ord_additional_contractor_external_id=contractor_ext_id)
    await state.set_state("vk_ord_additional_subject")
    await message.answer(
        "🧾 *Создание доп. соглашения (шаг 3/4)*\n\n"
        "Кратко опишите предмет договора.\n"
        "Например: *распространение рекламы на площадках VK*.",
        reply_markup=step_kb(),
        parse_mode="Markdown",
    )


async def vk_ord_additional_subject_step(message: _Message_vk, state: _FSMContext_vk):
    """
    Шаг 3/4: предмет доп. соглашения (человекочитаемый текст).
    """
    subject_text = (message.text or "").strip()
    if not subject_text:
        await message.answer(
            "Предмет договора не должен быть пустым. Опишите его или вернитесь назад.",
            reply_markup=step_kb(),
        )
        return

    await state.update_data(vk_ord_additional_subject_text=subject_text)
    await state.set_state("vk_ord_additional_date")
    await message.answer(
        "🧾 *Создание доп. соглашения (шаг 4/4)*\n\n"
        "Укажите дату заключения доп. соглашения в формате ДД.ММ.ГГГГ.",
        reply_markup=step_kb(),
        parse_mode="Markdown",
    )


async def vk_ord_additional_date_step(message: _Message_vk, state: _FSMContext_vk):
    """
    Шаг 4/4: дата заключения доп. соглашения.
    """
    date_raw = (message.text or "").strip()
    if not date_raw:
        await message.answer(
            "Дата не должна быть пустой. Укажите дату в формате ДД.ММ.ГГГГ.",
            reply_markup=step_kb(),
        )
        return

    await state.update_data(vk_ord_additional_date_raw=date_raw)
    data = await state.get_data()
    user_id = str(message.from_user.id)
    last_contract = _get_last_contract(user_id)

    parent_info = ""
    if last_contract:
        parent_info = (
            f"• Базовый договор (external_id): `{last_contract.get('external_id', '')}`\n"
            f"• Номер базового договора: `{last_contract.get('number', '')}`\n"
            f"• Дата базового договора: `{last_contract.get('date', '')}`\n"
        )
    else:
        parent_info = (
            "• Базовый договор пока не найден в данных бота.\n"
            "  Для создания доп. соглашения потребуется ранее созданный основной договор.\n"
        )

    text = (
        "Проверьте данные доп. соглашения:\n"
        f"• Заказчик (client_external_id): `{data.get('vk_ord_additional_client_external_id', '')}`\n"
        f"• Исполнитель (contractor_external_id): `{data.get('vk_ord_additional_contractor_external_id', '')}`\n"
        f"• Предмет: {data.get('vk_ord_additional_subject_text', '')}\n"
        f"• Дата заключения: `{date_raw}`\n"
        "\n"
        + parent_info +
        "\nЕсли всё верно — нажмите «✅ Подтвердить»."
    )

    kb = _ReplyKeyboardMarkup_vk(
        keyboard=[
            [_KeyboardButton_vk(text="✅ Подтвердить")],
            [_KeyboardButton_vk(text="◀  Назад"), _KeyboardButton_vk(text="✖  На главную")],
        ],
        resize_keyboard=True,
    )

    await state.set_state("vk_ord_additional_confirm")
    await message.answer(text, reply_markup=kb, parse_mode="Markdown")


async def vk_ord_additional_confirm_step(message: _Message_vk, state: _FSMContext_vk):
    """
    Финальное создание доп. соглашения типа additional через VK.ОРД API.
    Собирает данные из шага 1–4 и вызывает метод /v1/contract/{external_id}.
    """
    if (message.text or "").strip() != "✅ Подтвердить":
        await message.answer(
            "Чтобы создать доп. соглашение, нажмите «✅ Подтвердить» или используйте «◀  Назад»/«✖  На главную».",
            reply_markup=step_kb(),
        )
        return

    data = await state.get_data()
    user_id = str(message.from_user.id)
    ext_id = f"tg-{user_id}-additional-{int(_time_vk.time())}"

    last_contract = _get_last_contract(user_id)
    if not last_contract:
        await state.clear()
        await message.answer(
            "Не найден базовый договор для доп. соглашения.\n"
            "Сначала создайте основной договор через «📄 Добавить договор», а затем повторите попытку.",
            reply_markup=vk_ord_menu_kb(),
        )
        return

    client_ext_id = data.get("vk_ord_additional_client_external_id")
    contractor_ext_id = data.get("vk_ord_additional_contractor_external_id")
    subject_text = data.get("vk_ord_additional_subject_text", "")
    date_raw = data.get("vk_ord_additional_date_raw", "")

    # Определяем subject_type по тексту, по умолчанию distribution
    subj_low = (subject_text or "").lower()
    if "организац" in subj_low or "орг" in subj_low:
        subject_type = "org_distribution"
    elif "услуг" in subj_low:
        subject_type = "service"
    else:
        subject_type = "distribution"

    # Пробуем привести дату к формату ГГГГ-ММ-ДД для API
    date_norm = date_raw.replace("/", ".").replace("-", ".")
    parts = date_norm.split(".")
    date_api = date_raw
    if len(parts) == 3 and all(p.isdigit() for p in parts):
        dd, mm, yy = parts
        if len(yy) == 2:
            yy = "20" + yy
        date_api = f"{yy.zfill(4)}-{mm.zfill(2)}-{dd.zfill(2)}"

    parent_external_id = last_contract.get("external_id")

    payload = {
        "type": "additional",
        "client_external_id": client_ext_id,
        "contractor_external_id": contractor_ext_id,
        "date": date_api,
        # Можно использовать любое удобное обозначение серии, по умолчанию "1"
        "serial": data.get("vk_ord_additional_serial", "1"),
        "subject_type": subject_type,
        "flags": [
            "contractor_is_creatives_reporter",
        ],
        "parent_contract_external_id": parent_external_id,
        # Сумма доп. соглашения: по умолчанию 0, можно расширить мастером позже
        "amount": data.get("vk_ord_additional_amount", "0"),
    }

    ok, resp = await vk_ord_api_request(user_id, "PUT", f"/v1/contract/{ext_id}", payload)
    if not ok:
        await state.clear()
        await message.answer(
            "❌ Не удалось создать доп. соглашение через VK.ОРД API.\n\n"
            f"*Ответ сервера:* `{resp}`\n\n"
            "Сверьтесь с документацией VK.ОРД по методу создания доп. соглашения и скорректируйте данные.",
            parse_mode="Markdown",
            reply_markup=vk_ord_menu_kb(),
        )
        return

    text = "✅ Доп. соглашение успешно создано в VK.ОРД.\n"
    if isinstance(resp, dict):
        vk_id = resp.get("id")
        if vk_id:
            text += f"ID в VK.ОРД: `{vk_id}`\n"
    text += f"External ID (ваш): `{ext_id}`"

    await state.clear()
    await message.answer(text, parse_mode="Markdown", reply_markup=vk_ord_menu_kb())



async def vk_ord_service_serial_step(message: _Message_vk, state: _FSMContext_vk):
    """
    Шаг 1/7: серийный номер договора.
    """
    serial = (message.text or "").strip()
    if not serial:
        await message.answer(
            "Серийный номер договора не должен быть пустым. Укажите номер или вернитесь назад.",
            reply_markup=step_kb(),
        )
        return

    await state.update_data(vk_ord_service_serial=serial)
    await state.set_state("vk_ord_service_comment")
    await message.answer(
        "🧾 *Создание договора (Оказание услуг, шаг 2/7)*\n\n"
        "Укажите комментарий к договору (например, внутреннее описание или назначение).",
        reply_markup=step_kb(),
        parse_mode="Markdown",
    )


async def vk_ord_service_comment_step(message: _Message_vk, state: _FSMContext_vk):
    """
    Шаг 2/7: комментарий к договору.
    """
    comment = (message.text or "").strip()
    if not comment:
        await message.answer(
            "Комментарий к договору не должен быть пустым. Укажите текст или вернитесь назад.",
            reply_markup=step_kb(),
        )
        return

    await state.update_data(vk_ord_service_comment=comment)
    await state.set_state("vk_ord_service_client")
    await message.answer(
        "🧾 *Создание договора (Оказание услуг, шаг 3/7)*\n\n"
        "Укажите `external_id` заказчика (`client_external_id`).",
        reply_markup=step_kb(),
        parse_mode="Markdown",
    )


async def vk_ord_service_client_step(message: _Message_vk, state: _FSMContext_vk):
    """
    Шаг 3/7: заказчик. Пользователь вводит название контрагента или ИНН,
    бот ищет его в своём справочнике и подставляет external_id.
    """
    user_id = str(message.from_user.id)
    query = (message.text or "").strip()

    ext_id, person = _find_person_external_id(user_id, query)
    if not ext_id:
        await message.answer(
            "Не удалось найти контрагента с таким названием или ИНН.\n"
            "Сначала создайте контрагента через «➕ Добавить контрагента», "
            "а затем повторите ввод заказчика.",
            reply_markup=step_kb(),
        )
        return


    ext_id, person = _find_person_external_id(user_id, query)
    if not ext_id:
        await message.answer(
            "Не удалось найти контрагента с таким названием или ИНН.\n"
            "Сначала создайте контрагента через «➕ Внести контрагента», "
            "а затем повторите ввод заказчика.",
            reply_markup=step_kb(),
        )
        return

    await state.update_data(
        vk_ord_service_client_external_id=ext_id,
        vk_ord_service_client_name=person.get("name"),
    )
    await state.set_state("vk_ord_service_contractor")
    await message.answer(
        "🧾 *Создание договора (Оказание услуг, шаг 4/7)*\n\n"
        "Укажите `external_id` исполнителя (он же издатель) — `contractor_external_id`.",
        reply_markup=step_kb(),
        parse_mode="Markdown",
    )


async def vk_ord_service_contractor_step(message: _Message_vk, state: _FSMContext_vk):
    """
    Шаг 4/7: исполнитель (contractor_external_id).
    Пользователь вводит название контрагента или ИНН,
    бот ищет его в своём справочнике и подставляет external_id.
    """
    user_id = str(message.from_user.id)
    query = (message.text or "").strip()

    ext_id, person = _find_person_external_id(user_id, query)
    if not ext_id:
        await message.answer(
            "Не удалось найти исполнителя с таким названием или ИНН.\n"
            "Сначала создайте этого контрагента через «➕ Добавить контрагента», "
            "а затем повторите ввод исполнителя.",
            reply_markup=step_kb(),
        )
        return


    ext_id, person = _find_person_external_id(user_id, query)
    if not ext_id:
        await message.answer(
            "Не удалось найти исполнителя с таким названием или ИНН.\n"
            "Сначала создайте этого контрагента через «➕ Добавить контрагента», "
            "а затем повторите ввод исполнителя.",
            reply_markup=step_kb(),
        )
        return


    await state.update_data(
        vk_ord_service_contractor_external_id=ext_id,
        vk_ord_service_contractor_name=person.get("name"),
    )
    await state.set_state("vk_ord_service_subject")
    kb = _ReplyKeyboardMarkup_vk(
        keyboard=[
            [
                _KeyboardButton_vk(text="Посредничество"),
                _KeyboardButton_vk(text="Представительство"),
            ],
            [
                _KeyboardButton_vk(text="Организация распространения"),
                _KeyboardButton_vk(text="Распространение рекламы"),
            ],
            [
                _KeyboardButton_vk(text="Иное"),
            ],
            [
                _KeyboardButton_vk(text="◀  Назад"),
                _KeyboardButton_vk(text="✖  На главную"),
            ],
        ],
        resize_keyboard=True,
    )
    await message.answer(
        "🧾 *Создание договора (Оказание услуг, шаг 5/7)*\n\n"
        "Выберите предмет договора из списка:\n"
        "• Посредничество\n"
        "• Представительство\n"
        "• Организация распространения\n"
        "• Распространение рекламы\n"
        "• Иное",
        reply_markup=kb,
        parse_mode="Markdown",
    )


def _vk_ord_map_service_subject(text: str) -> tuple[str, str]:
    """
    Маппинг человекочитаемого предмета договора на код subject_type VK.ОРД.
    Возвращает пару: (код subject_type, человекочитаемое название).

    Логика:
    - "Посредничество"      -> service
    - "Представительство"   -> representation
    - "Распространение рекламы" -> distribution
    - "Организация распространения" -> org_distribution
    - всё остальное         -> other
    """
    t = (text or "").strip().lower()

    if "посред" in t:
        return "service", "Посредничество"
    if "представ" in t:
        return "representation", "Представительство"
    # Сначала отдельно обрабатываем "распространение рекламы"
    if "реклам" in t:
        return "distribution", "Распространение рекламы"
    # "Организация распространения" — только если явно есть и "организац", и "распростран"
    if "организац" in t and "распростран" in t:
        return "org_distribution", "Организация распространения"

    return "other", "Иное"



async def vk_ord_service_subject_step(message: _Message_vk, state: _FSMContext_vk):
    """
    Шаг 5/7: предмет договора (subject_type).
    """
    subj_raw = (message.text or "").strip()
    subject_type, subject_human = _vk_ord_map_service_subject(subj_raw)

    await state.update_data(
        vk_ord_service_subject_type=subject_type,
        vk_ord_service_subject_human=subject_human,
    )
    await state.set_state("vk_ord_service_date")
    await message.answer(
        "🧾 *Создание договора (Оказание услуг, шаг 6/7)*\n\n"
        "Укажите дату заключения договора в формате ДД.ММ.ГГГГ.\n"
        "Если даты нет — отправьте пустое сообщение или напишите «нет».",
        reply_markup=step_kb(),
        parse_mode="Markdown",
    )


async def vk_ord_service_date_step(message: _Message_vk, state: _FSMContext_vk):
    """
    Шаг 6/7: дата заключения договора (может быть пустой).
    Проверяем, что дата не раньше 01.01.1991 и не позже сегодняшнего дня,
    как того требует VK.ОРД (date_constraint).
    """
    raw = (message.text or "").strip()

    # Пустая дата или "нет" — разрешаем и считаем, что даты нет
    if not raw or raw.lower() == "нет":
        date_raw = ""
    else:
        # Нормализуем разделители и пробуем распарсить ДД.ММ.ГГГГ или ДД.ММ.ГГ
        norm = raw.replace("/", ".").replace("-", ".")
        parts = [p for p in norm.split(".") if p]
        if len(parts) != 3 or not all(p.isdigit() for p in parts):
            await message.answer(
                "Дата должна быть в формате ДД.ММ.ГГГГ. Попробуйте ещё раз.",
                reply_markup=step_kb(),
            )
            return

        dd, mm, yy = parts
        if len(yy) == 2:
            yy = "20" + yy
        try:
            d = datetime.date(int(yy), int(mm), int(dd))
        except ValueError:
            await message.answer(
                "Не удалось распознать дату. Убедитесь, что она существует в календаре (ДД.ММ.ГГГГ).",
                reply_markup=step_kb(),
            )
            return

        min_date = datetime.date(1991, 1, 1)
        today = now_tz().date()
        if d < min_date or d > today:
            await message.answer(
                "Дата договора не может быть раньше 01.01.1991 и позже сегодняшнего дня.\n"
                f"Вы указали: {d.strftime('%d.%m.%Y')}. Попробуйте ещё раз.",
                reply_markup=step_kb(),
            )
            return

        # Сохраняем строку в привычном для пользователя формате
        date_raw = d.strftime("%d.%m.%Y")

    await state.update_data(vk_ord_service_date_raw=date_raw)
    await state.set_state("vk_ord_service_amount")
    await message.answer(
        "🧾 *Создание договора (Оказание услуг, шаг 7/7)*\n\n"
        "Укажите сумму договора в рублях (без знака ₽).\n"
        "Если сумма не фиксирована — отправьте 0 или напишите «нет».",
        reply_markup=step_kb(),
        parse_mode="Markdown",
    )


async def vk_ord_service_amount_step(message: _Message_vk, state: _FSMContext_vk):
    """
    Ввод суммы договора и показ финального подтверждения.
    """
    raw = (message.text or "").strip().replace(" ", "").replace(",", ".")
    if not raw:
        raw = "0"
    if raw.lower() in {"нет", "неизвестно"}:
        raw = "0"

    # Простейшая проверка числа
    try:
        float(raw)
    except ValueError:
        await message.answer(
            "Сумма договора должна быть числом. Попробуйте ещё раз.",
            reply_markup=step_kb(),
        )
        return

    await state.update_data(vk_ord_service_amount_raw=raw)
    data = await state.get_data()
    user_id = str(message.from_user.id)

    # Нормализуем дату к ГГГГ-ММ-ДД при возможности
    date_raw = data.get("vk_ord_service_date_raw", "") or ""
    date_norm = date_raw.replace("/", ".").replace("-", ".")
    parts = [p for p in date_norm.split(".") if p]
    date_api = ""
    if len(parts) == 3 and all(p.isdigit() for p in parts):
        dd, mm, yy = parts
        if len(yy) == 2:
            yy = "20" + yy
        date_api = f"{yy.zfill(4)}-{mm.zfill(2)}-{dd.zfill(2)}"

    if not date_api:
        date_api = ""

    text = (
        "Проверьте данные договора (Оказание услуг):\n"
        f"• Серийный номер: `{data.get('vk_ord_service_serial', '')}`\n"
        f"• Комментарий: {data.get('vk_ord_service_comment', '')}\n"
        f"• Заказчик: {data.get('vk_ord_service_client_name', '')} "
        f"(client_external_id: `{data.get('vk_ord_service_client_external_id', '')}`)\n"
        f"• Исполнитель: {data.get('vk_ord_service_contractor_name', '')} "
        f"(contractor_external_id: `{data.get('vk_ord_service_contractor_external_id', '')}`)\n"
        f"• Предмет: {data.get('vk_ord_service_subject_human', '')} "
        f"(subject_type: `{data.get('vk_ord_service_subject_type', '')}`)\n"
        f"• Дата заключения: `{date_api or date_raw}`\n"
        f"• Сумма: {raw} руб.\n"
        "\nЕсли всё верно — нажмите «✅ Подтвердить»."
    )

    kb = _ReplyKeyboardMarkup_vk(
        keyboard=[
            [_KeyboardButton_vk(text="✅ Подтвердить")],
            [_KeyboardButton_vk(text="◀  Назад"), _KeyboardButton_vk(text="✖  На главную")],
        ],
        resize_keyboard=True,
    )

    await state.set_state("vk_ord_service_confirm")
    await message.answer(text, reply_markup=kb, parse_mode=None)


async def vk_ord_service_confirm_step(message: _Message_vk, state: _FSMContext_vk):
    """
    Финальное создание договора типа service (Оказание услуг) через VK.ОРД API.
    """
    if (message.text or "").strip() != "✅ Подтвердить":
        await message.answer(
            "Чтобы создать договор, нажмите «✅ Подтвердить» или используйте «◀  Назад»/«✖  На главную».",
            reply_markup=step_kb(),
        )
        return

    data = await state.get_data()
    user_id = str(message.from_user.id)
    ext_id = f"tg-{user_id}-service-{int(_time_vk.time())}"

    client_ext_id = data.get("vk_ord_service_client_external_id")
    contractor_ext_id = data.get("vk_ord_service_contractor_external_id")
    subject_type = data.get("vk_ord_service_subject_type", "service")
    serial = data.get("vk_ord_service_serial", "")
    amount_raw = data.get("vk_ord_service_amount_raw", "0")

    # Нормализуем дату для отправки в API
    date_raw = data.get("vk_ord_service_date_raw", "") or ""
    date_norm = date_raw.replace("/", ".").replace("-", ".")
    parts = [p for p in date_norm.split(".") if p]
    date_api = ""
    if len(parts) == 3 and all(p.isdigit() for p in parts):
        dd, mm, yy = parts
        if len(yy) == 2:
            yy = "20" + yy
        date_api = f"{yy.zfill(4)}-{mm.zfill(2)}-{dd.zfill(2)}"
    if not date_api and date_raw.strip().lower() in {"", "нет"}:
        date_api = ""
    elif not date_api:
        # Если не смогли распарсить, отправим как есть
        date_api = date_raw

    payload = {
        "type": "service",
        "client_external_id": client_ext_id,
        "contractor_external_id": contractor_ext_id,
        "date": date_api,
        "serial": serial,
        "subject_type": subject_type,
        "flags": [
            "vat_included",
            "contractor_is_creatives_reporter",
        ],
        "amount": amount_raw,
    }

    ok, resp = await vk_ord_api_request(user_id, "PUT", f"/v1/contract/{ext_id}", payload)
    if not ok:
        await state.clear()
        await message.answer(
            "❌ Не удалось создать договор (Оказание услуг) через VK.ОРД API.\n\n"
            f"*Ответ сервера:* `{resp}`\n\n"
            "Сверьтесь с документацией VK.ОРД по методу создания договора оказания услуг и скорректируйте данные.",
            parse_mode="Markdown",
            reply_markup=vk_ord_menu_kb(),
        )
        return

    # Сохраняем договор как последний, чтобы можно было создавать доп. соглашения
    _set_last_contract(
        user_id,
        ext_id,
        serial,
        date_api or date_raw,
    )

    text = "✅ Договор (Оказание услуг) успешно создан в VK.ОРД.\n"
    if isinstance(resp, dict):
        vk_id = resp.get("id")
        if vk_id:
            text += f"ID в VK.ОРД: `{vk_id}`\n"
    text += f"External ID (ваш): `{ext_id}`"

    await state.clear()
    await message.answer(text, parse_mode="Markdown", reply_markup=vk_ord_menu_kb())


async def vk_ord_contract_number_step(message: _Message_vk, state: _FSMContext_vk):
    num = (message.text or "").strip()
    await state.update_data(vk_ord_contract_number=num)
    await state.set_state("vk_ord_contract_date")
    await message.answer(
        "📄 *Создание договора (шаг 2/4)*\n\n"
        "Укажите дату договора в формате ДД.ММ.ГГГГ.",
        reply_markup=step_kb(),
        parse_mode="Markdown",
    )


async def vk_ord_contract_date_step(message: _Message_vk, state: _FSMContext_vk):
    date_str = (message.text or "").strip()
    await state.update_data(vk_ord_contract_date=date_str)
    await state.set_state("vk_ord_contract_subject")
    await message.answer(
        "📄 *Создание договора (шаг 3/4)*\n\n"
        "Кратко опишите предмет договора.\n"
        "Например: «оказание услуг по размещению рекламной информации в Telegram-канале».",
        reply_markup=step_kb(),
        parse_mode="Markdown",
    )


async def vk_ord_contract_subject_step(message: _Message_vk, state: _FSMContext_vk):
    subj = (message.text or "").strip()
    await state.update_data(vk_ord_contract_subject=subj)
    await state.set_state("vk_ord_contract_amount")
    await message.answer(
        "📄 *Создание договора (шаг 4/4)*\n\n"
        "Укажите общую сумму договора в рублях (без знака ₽).\n"
        "Если сумма не фиксирована — отправьте 0 или напишите «нет».",
        reply_markup=step_kb(),
        parse_mode="Markdown",
    )


async def vk_ord_contract_amount_step(message: _Message_vk, state: _FSMContext_vk):
    raw = (message.text or "").strip().replace(" ", "").replace(",", ".")
    if raw.lower() == "нет":
        raw = "0"
    await state.update_data(vk_ord_contract_amount_raw=raw)

    data = await state.get_data()
    last_person = _get_last_person(str(message.from_user.id)) or {}
    text = (
        "Проверьте данные договора:\n"
        f"• Номер: *{data.get('vk_ord_contract_number', '')}*\n"
        f"• Дата: `{data.get('vk_ord_contract_date', '')}`\n"
        f"• Предмет: {data.get('vk_ord_contract_subject', '')}\n"
        f"• Сумма: {raw} руб.\n"
        f"• Контрагент (advertiser): {last_person.get('name', '—')} (ext_id: {last_person.get('external_id', '—')})\n\n"
        "Если всё верно — нажмите «✅ Подтвердить»."
    )
    kb = _ReplyKeyboardMarkup_vk(
        keyboard=[
            [_KeyboardButton_vk(text="✅ Подтвердить")],
            [_KeyboardButton_vk(text="◀  Назад"), _KeyboardButton_vk(text="✖  На главную")],
        ],
        resize_keyboard=True,
    )
    await state.set_state("vk_ord_contract_confirm")
    await message.answer(text, reply_markup=kb, parse_mode="Markdown")


async def vk_ord_contract_confirm_step(message: _Message_vk, state: _FSMContext_vk):
    if (message.text or "").strip() != "✅ Подтвердить":
        await message.answer(
            "Чтобы создать договор, нажмите «✅ Подтвердить» или используйте «◀  Назад»/«✖  На главную».",
            reply_markup=step_kb(),
        )
        return

    data = await state.get_data()
    user_id = str(message.from_user.id)
    ext_id = f"tg-{user_id}-contract-{int(_time_vk.time())}"

    payload = {
        "external_id": ext_id,
        "number": data.get("vk_ord_contract_number"),
        "date": data.get("vk_ord_contract_date"),
        "subject": data.get("vk_ord_contract_subject"),
        "amount": data.get("vk_ord_contract_amount_raw") or "0",
        "type": "main",
        "persons": [
            {
                "person_external_id": data.get("vk_ord_contract_person_external_id"),
                "role": "advertiser",
            }
        ],
    }

    ok, resp = await vk_ord_api_request(user_id, "PUT", f"/v1/contract/{ext_id}", payload)
    if not ok:
        await message.answer(
            "❌ Не удалось создать договор через VK.ОРД API.\n\n"
            f"*Ответ сервера:*\n`{resp}`",
            parse_mode="Markdown",
            reply_markup=vk_ord_menu_kb(),
        )
        await state.clear()
        return

    _set_last_contract(
        user_id,
        ext_id,
        data.get("vk_ord_contract_number", ""),
        data.get("vk_ord_contract_date", ""),
    )
    text = "✅ Договор успешно создан в VK.ОРД.\n"
    if isinstance(resp, dict):
        vk_id = resp.get("id")
        if vk_id:
            text += f"ID в VK.ОРД: `{vk_id}`\n"
    text += f"External ID (ваш): `{ext_id}`"
    await message.answer(text, parse_mode="Markdown", reply_markup=vk_ord_menu_kb())
    await state.clear()
# ---------- МАСТЕР СОЗДАНИЯ КРЕАТИВА / ERID ----------

async def vk_ord_add_creative(message: _Message_vk, state: _FSMContext_vk):
    user_id = str(message.from_user.id)
    last_contract = _get_last_contract(user_id)
    last_person = _get_last_person(user_id)

    if not last_contract and not last_person:
        await message.answer(
            "Чтобы оформить креатив и получить ERID, сначала нужно:\n"
            "1) Создать контрагента (через «➕ Добавить контрагента»)\n"
            "2) Создать договор (через «🖥️ Отправить договор в ЕРИР»)\n\n"
            "После этого вернитесь к оформлению креатива.",
            reply_markup=vk_ord_menu_kb(),
        )
        return

    # дальше — остальной код функции без изменения отступов


    last_contract = _get_last_contract(user_id)
    last_person = _get_last_person(user_id)
    if not last_contract and not last_person:
        await message.answer(
            "Чтобы оформить креатив и получить ERID, сначала нужно:\n"
            "1) Создать контрагента (через «➕ Добавить контрагента»)\n"
            "2) Создать договор (через «🖥️ Отправить договор в ЕРИР»)\n\n"
            "После этого вернитесь к оформлению креатива.",
            reply_markup=vk_ord_menu_kb(),
        )
        return


    await state.clear()
    if last_contract:
        await state.update_data(vk_ord_creative_contract_external_id=last_contract["external_id"])
    elif last_person:
        await state.update_data(vk_ord_creative_person_external_id=last_person["external_id"])

    await state.set_state("vk_ord_creative_name")
    await message.answer(
        "🎨 *Оформление креатива (шаг 1/5)*\n\n"
        "Укажите краткое название креатива.\n"
        "Например: «Реклама бота PromoPro в ТГ».",
        reply_markup=step_kb(),
        parse_mode="Markdown",
    )


async def vk_ord_creative_name_step(message: _Message_vk, state: _FSMContext_vk):
    name = (message.text or "").strip()
    await state.update_data(vk_ord_creative_name=name)
    await state.set_state("vk_ord_creative_url")
    await message.answer(
        "🎨 *Оформление креатива (шаг 2/5)*\n\n"
        "Укажите ссылку (или несколько ссылок), где будет размещён креатив.\n"
        "Если ссылок несколько, укажите их через пробел или с новой строки.",
        reply_markup=step_kb(),
        parse_mode="Markdown",
    )


async def vk_ord_creative_url_step(message: _Message_vk, state: _FSMContext_vk):
    url = (message.text or "").strip()
    await state.update_data(vk_ord_creative_url=url)
    await state.set_state("vk_ord_creative_period")
    await message.answer(
        "🎨 *Оформление креатива (шаг 3/5)*\n\n"
        "Укажите *период размещения* креатива.\n"
        "Формат: *ДД.ММ.ГГГГ–ДД.ММ.ГГГГ (через тире)*.\n\n"
        "_Например: 15.11.2025–20.11.2025_",
        reply_markup=step_kb(),
        parse_mode="Markdown",
    )


async def vk_ord_creative_period_step(message: _Message_vk, state: _FSMContext_vk):
    period = (message.text or "").strip()
    await state.update_data(vk_ord_creative_period=period)
    await state.set_state("vk_ord_creative_texts")
    await message.answer(
        "🎨 *Оформление креатива (шаг 4/5)*\n\n"
        "Введите текстовые данные креатива.\n"
        "Например:\n\n"
        "_Хочешь маркировать посты без хлопот?_\n"
        "_Переходи на сторону PROMO-PRO_\n"
        "_У нас есть печеньки.._",
        reply_markup=step_kb(),
        parse_mode="Markdown",
    )


async def vk_ord_creative_texts_step(message: _Message_vk, state: _FSMContext_vk):
    texts_raw = (message.text or "").strip()
    await state.update_data(vk_ord_creative_texts_raw=texts_raw)
    await state.set_state("vk_ord_creative_media")
    await message.answer(
        "🎨 *Оформление креатива (шаг 5/6)*\n\n"
        "Отправьте картинку/видео/файл с креативом одним сообщением — "
        "бот автоматически подгрузит медиа в VK.ОРД и прикрепит к данному креативу.\n",
        reply_markup=step_kb(),
        parse_mode="Markdown",
    )


async def vk_ord_creative_media_step(message: _Message_vk, state: _FSMContext_vk):
    """
    Шаг 5 мастера креатива:
    • если пользователь отправил медиафайл — загружаем его в VK.ОРД
      и сохраняем полученный external_id;
    • если прислал текст — трактуем как external_id(ы), как раньше.
    """
    user = message.from_user
    user_id = str(user.id) if user else "0"

    # 1. Пытаемся взять медиа из сообщения Telegram
    media_info = await _vk_ord_extract_telegram_media(message)
    if media_info is not None:
        file_bytes, filename, content_type = media_info

        await message.answer("⏳ Загружаю медиафайл в VK.ОРД…")

        ok, result = await vk_ord_upload_media(
            user_id=user_id,
            file_bytes=file_bytes,
            filename=filename,
            content_type=content_type,
        )

        if not ok:
            # Ошибка загрузки — даём пользователю текст ошибки и остаёмся на этом же шаге
            await message.answer(
                "❌ Не удалось загрузить медиафайл в VK.ОРД.\n"
                "Вы можете повторить попытку и отправить файл ещё раз "
                "или указать `external_id` файла вручную.\n\n"
                f"Технические детали: `{str(result)[:500]}`",
                reply_markup=step_kb(),
                parse_mode="Markdown",
            )
            await state.set_state("vk_ord_creative_media")
            return

        external_id = str(result)
        # Сохраняем external_id как будто его ввёл пользователь
        await state.update_data(vk_ord_creative_media_raw=external_id)

        # Переходим к шагу ККТУ
        await state.set_state("vk_ord_creative_kktu")
        await message.answer(
            "✅ Файл загружен в VK.ОРД.\n"
            f"Его `external_id`: `{external_id}`.\n\n"
            "Теперь укажите код(ы) ККТУ рекламируемого товара или услуги.\n"
            "Можно несколько через запятую. Если не знаете — напишите «нет».",
            reply_markup=step_kb(),
            parse_mode="Markdown",
        )
        return

    # 2. Если медиа нет — оставляем старое поведение: пользователь вводит external_id руками
    media_raw = (message.text or "").strip()
    if not media_raw:
        await message.answer(
            "Отправьте картинку/видео/файл с креативом или укажите `external_id` "
            "уже загруженного в VK.ОРД медиафайла.",
            reply_markup=step_kb(),
            parse_mode="Markdown",
        )
        await state.set_state("vk_ord_creative_media")
        return

    await state.update_data(vk_ord_creative_media_raw=media_raw)
    await state.set_state("vk_ord_creative_kktu")
    await message.answer(
        "🎨 *Оформление креатива (шаг 6/6)*\n\n"
        "Укажите код(ы) ККТУ рекламируемого товара или услуги.\n"
        "Можно несколько через запятую. Если не знаете — напишите «нет».",
        reply_markup=step_kb(),
        parse_mode="Markdown",
    )



async def vk_ord_creative_kktu_step(message: _Message_vk, state: _FSMContext_vk):
    kktu_raw = (message.text or "").strip()
    await state.update_data(vk_ord_creative_kktu_raw=kktu_raw)

    data = await state.get_data()
    texts_raw = (data.get("vk_ord_creative_texts_raw") or "").strip()

    # Сформируем короткий превью текстов для экрана подтверждения
    texts_preview = ""
    if texts_raw:
        _lines = [ln.strip() for ln in texts_raw.splitlines() if ln.strip()]
        if _lines:
            if len(_lines) == 1:
                texts_preview = _lines[0]
            else:
                texts_preview = "; ".join(_lines[:3])
                if len(_lines) > 3:
                    texts_preview += " …"

    media_raw = (data.get("vk_ord_creative_media_raw") or "").strip()

    text = (
        "Проверьте данные креатива:\n"
        f"• Название: *{data.get('vk_ord_creative_name', '')}*\n"
        f"• URL: `{data.get('vk_ord_creative_url', '')}`\n"
        f"• Период: {data.get('vk_ord_creative_period', '')}\n"
        f"• Тексты: {texts_preview or 'не заданы'}\n"
        f"• Медиа `external_id`: {media_raw or 'не указаны'}\n"
        f"• ККТУ: {kktu_raw or 'не указаны'}\n\n"
        "Если всё верно — нажмите «✅ Подтвердить»."
    )
    kb = _ReplyKeyboardMarkup_vk(
        keyboard=[
            [_KeyboardButton_vk(text="✅ Подтвердить")],
            [_KeyboardButton_vk(text="◀  Назад"), _KeyboardButton_vk(text="✖  На главную")],
        ],
        resize_keyboard=True,
    )
    await state.set_state("vk_ord_creative_confirm")
    await message.answer(text, reply_markup=kb, parse_mode="Markdown")
async def vk_ord_creative_confirm_step(message: _Message_vk, state: _FSMContext_vk):
    """
    Финальный шаг мастера создания креатива (ERID) через VK.ОРД.

    Использует метод v3:
      PUT /v3/creative/{external_id}

    external_id генерируется на стороне бота.
    Креатив привязывается к "последнему договору" (_get_last_contract).
    """
    user = message.from_user
    user_id = str(user.id) if user else "0"

    # Проверяем, авторизован ли пользователь в VK.ОРД
    if not user_is_authorized(user_id):
        await message.answer(
            "Для оформления креатива сначала подключите кабинет VK.ОРД через кнопку "
            "«➦ Перейти в кабинет «VK.ОРД»».",
            reply_markup=main_kb()
        )
        await state.clear()
        return

    data = await state.get_data()

    # Данные, собранные на предыдущих шагах мастера
    name_raw = (data.get("vk_ord_creative_name") or "").strip()
    url_raw = (data.get("vk_ord_creative_url") or "").strip()
    period_raw = (data.get("vk_ord_creative_period") or "").strip()
    texts_raw = (data.get("vk_ord_creative_texts_raw") or "").strip()
    media_raw = (data.get("vk_ord_creative_media_raw") or "").strip()
    kktu_raw = (data.get("vk_ord_creative_kktu_raw") or "").strip()

    # Нормализуем поля для Markdown
    name_md = md_escape(name_raw or "Без названия")
    url_md = md_escape(url_raw or "—")
    period_md = md_escape(period_raw or "—")

    # KKTU — список строк, разделённых пробелами/запятыми/переводами строк
    kktus = []
    if kktu_raw:
        # Разбиваем по любым пробелам/переводам строки/табам, игнорируя запятые
        parts = _re_vk.split(r"\s+", kktu_raw.replace(",", " ").strip())
        cleaned = []
        for p in parts:
            p = (p or "").strip()
            if not p:
                continue
            # Отбрасываем заведомо неверные значения (оставляем только цифры и точки)
            if not _re_vk.fullmatch(r"[0-9.]+", p):
                continue
            cleaned.append(p)
        kktus = cleaned

    if not kktus:
        await message.answer(
            "⚠️ Не удалось распознать ни одного кода ККТУ.\n"
            "Проверьте ввод и попробуйте ещё раз.",
            reply_markup=step_kb()
        )
        await state.set_state("vk_ord_creative_kktu")
        return

    # Привязка к последнему договору VK.ОРД
    last_contract = _get_last_contract(user_id)
    if not last_contract or not last_contract.get("external_id"):
        await message.answer(
            "❌ Не найден последний договор в VK.ОРД.\n\n"
            "Сначала создайте договор через «📄 Добавить договор», "
            "а затем повторите оформление креатива.",
            reply_markup=main_kb()
        )
        await state.clear()
        return

    contract_external_id = last_contract["external_id"]

    # Генерируем external_id креатива (можно любая уникальная строка)
    creative_external_id = f"cr-{int(_time_vk.time())}-{user_id}"

    # Подготовим тексты креатива из введённых строк
    texts = []
    if texts_raw:
        for _ln in texts_raw.splitlines():
            _ln = _ln.strip()
            if _ln:
                texts.append(_ln)
    if not texts:
        texts = [name_raw or "Рекламный баннер"]

    # Описание: берём первый текст, если он есть, иначе формируем по договору и периоду
    description_text = texts[0] if texts else ""
    if not description_text:
        description_text = (
            f"Реклама по договору {last_contract.get('number') or ''}. "
            f"Период: {period_raw or 'не указан'}."
        ).strip()

    
    # Подготовим media_external_ids: разбиваем по запятым/пробелам/переводам строк
    media_external_ids: list[str] = []
    if media_raw:
        parts = _re_vk.split(r"[\s,]+", media_raw.strip())
        for p in parts:
            p = (p or "").strip()
            if not p:
                continue
            media_external_ids.append(p)

    if not media_external_ids:
        await message.answer(
            "⚠️ Не удалось распознать ни одного media external_id.\n"
            "Укажите хотя бы один корректный идентификатор медиафайла VK.ОРД.",
            reply_markup=step_kb()
        )
        await state.set_state("vk_ord_creative_media")
        return

# Формируем тело запроса по примеру из документации v3/creative
    # https://sandbox.ord.vk.com/help/api/ref/creative.html
    body = {
        "contract_external_ids": [contract_external_id],
        "kktus": kktus,
        "name": name_raw or "Рекламный креатив",
        # Бренд и категория можно заполнять тем же, что и название/описание
        "brand": name_raw or "Без бренда",
        "category": "Рекламный баннер",
        "description": description_text,
        # Для простоты фиксируем тип оплаты/форму — как в примере
        "pay_type": "cpm",
        "form": "banner",
        # Таргетинг можно заполнить общей фразой или оставить как описательный текст
        "targeting": "Таргетинг не указан (создано через бота PROMO-PRO).",
        # Целевые URL — один или несколько URL, введённых пользователем.
        # Можно указать несколько ссылок через пробел, запятую или перенос строки.
        "target_urls": [u for u in _re_vk.split(r"[\s,]+", url_raw) if u] if url_raw else [],
        # Тексты креатива — список строк (предложений)
        "texts": texts,
        # media_external_ids — external_id(ы) медиафайлов, которые пользователь указал на шаге мастера.
        "media_external_ids": media_external_ids,
    }

    await message.answer("⏳ Отправляю данные креатива в VK.ОРД…")

    ok, resp = await vk_ord_api_request(
        user_id=user_id,
        method="PUT",
        path=["v3", "creative", creative_external_id],
        json_body=body,
    )

    if not ok:
        # Пытаемся красиво разобрать ошибку VK.ОРД
        human_msg = "❌ Не удалось создать креатив (ERID) через VK.ОРД API.\n\n"
        details = resp
        if isinstance(resp, dict):
            # Ошибка "creative_external_media_not_found" означает,
            # что VK.ОРД не нашёл ни одного медиафайла с указанным external_id.
            errors = resp.get("errors") or resp.get("error") or []
            if isinstance(errors, list):
                for err in errors:
                    if not isinstance(err, dict):
                        continue
                    code = err.get("error_code") or err.get("code")
                    msg = err.get("message") or ""
                    if code == "creative_external_media_not_found":
                        human_msg += (
                            "VK.ОРД не нашёл медиафайлы с указанным `external_id`.\n"
                            "Проверьте, что:\n"
                            "• файл действительно загружен в личный кабинет VK.ОРД;\n"
                            "• вы скопировали `external_id` именно этого файла без лишних символов;\n"
                            "• используемый кабинет (sandbox/prod) совпадает с тем, где был загружен файл.\n\n"
                        )
                        if msg:
                            human_msg += f"Сообщение VK.ОРД: {msg}\n\n"
                        break
            details = _json_vk.dumps(resp, ensure_ascii=False)
        else:
            details = str(resp)

        human_msg += f"Технические детали (для разработчика): {details}"
        await message.answer(human_msg, reply_markup=main_kb(), parse_mode="Markdown")
        await state.clear()
        return

    # Пытаемся достать ERID/ID из ответа VK.ОРД
    erid = resp.get("erid") or resp.get("id") or creative_external_id
    erid_md = md_escape(erid)

    text = (
        "✅ *Креатив успешно создан в VK.ОРД!*\n\n"
        f"• ERID: `{erid_md}`\n"
        f"• Название: *{name_md}*\n"
        f"• URL: {url_md}\n"
        f"• Период: {period_md}\n"
        f"• KKTU: {', '.join(kktus)}"
    )
    await message.answer(text, reply_markup=main_kb())
    await state.clear()

    erid = None
    if isinstance(resp, dict):
        erid = resp.get("erid")

    if erid:
        await message.answer(
            f"✅ Креатив зарегистрирован.\nERID: `{erid}`",
            parse_mode="Markdown",
            reply_markup=vk_ord_menu_kb(),
        )
    else:
        await message.answer(
            "✅ Креатив зарегистрирован через VK.ОРД API.\n"
            "Однако в ответе не найдено поле `erid`. Проверьте структуру ответа по swagger и при необходимости "
            "доработайте обработчик.",
            reply_markup=vk_ord_menu_kb(),
        )

    await state.clear()
# ================== ПОИСК ПО ИНН ====================
import os as _os_inn
from docx import Document as _Document_inn
from aiogram.types import Message as _Message_inn, InlineKeyboardMarkup as _InlineKeyboardMarkup_inn, InlineKeyboardButton as _InlineKeyboardButton_inn, CallbackQuery as _CallbackQuery_inn
from aiogram.fsm.context import FSMContext as _FSMContext_inn

GENERATED_PATH = r"D:\TRAFFIC\PROMO-PRO\generated"


def inn_pagination_kb() -> _InlineKeyboardMarkup_inn:
    return _InlineKeyboardMarkup_inn(
        inline_keyboard=[
            [
                _InlineKeyboardButton_inn(text="Назад", callback_data="inn_prev"),
                _InlineKeyboardButton_inn(text="Далее", callback_data="inn_next"),
            ],
            [_InlineKeyboardButton_inn(text="В главное меню", callback_data="inn_main")],
        ]
    )


def build_inn_summary_from_paragraphs(paragraphs, file_path: str, inn: str) -> str:
    lines = []
    for p in paragraphs:
        try:
            t = (p.text or "").strip()
        except Exception:
            t = ""
        if t:
            lines.append(t)
    if not lines:
        return f"Документ: {os.path.basename(file_path)}\nИНН: {inn}"

    header = ""
    for ln in lines:
        low = ln.lower()
        if ("счёт-оферта" in low or "счет-оферта" in low) and "№" in ln:
            header = ln.strip()
            break
    if not header:
        for ln in lines:
            low = ln.lower()
            if "договор" in low and "№" in ln:
                header = ln.strip()
                break
    if not header:
        header = f"Документ: {os.path.basename(file_path)}"
    if not header.startswith("🧾"):
        header = "🧾 " + header

    customer_line = ""
    for ln in lines:
        if ln.lower().startswith("заказчик"):
            customer_line = ln.strip()
            break

    inn_line = ""
    for ln in lines:
        if "инн" in ln.lower():
            inn_line = ln.strip()
            break
    if not inn_line:
        inn_line = f"ИНН: {inn}"

    ogrn_line = ""
    for ln in lines:
        if "огрн" in ln.lower():
            ogrn_line = ln.strip()
            break
    if not ogrn_line:
        ogrn_line = "ОГРН|ОГРНИП: —"

    period_line = ""
    for ln in lines:
        low = ln.lower()
        if "период" in low or "срок оказания услуг" in low or "срок размещения" in low:
            period_line = ln.strip()
            break

    count_line = ""
    for ln in lines:
        low = ln.lower()
        if "кол-во услуг" in low or "количество услуг" in low:
            count_line = ln.strip()
            break
    if count_line and not count_line.startswith("╰⪼"):
        count_line = "╰⪼" + count_line

    total_line = ""
    for ln in lines:
        if "общая сумма" in ln.lower():
            total_line = ln.strip()
            break
    if total_line and not total_line.startswith("💲"):
        total_line = "💲 " + total_line

    words_line = ""
    if total_line and total_line in lines:
        start_idx = lines.index(total_line)
        for ln in lines[start_idx + 1:]:
            if "руб" in ln.lower():
                words_line = ln.strip()
                break
    if not words_line:
        for ln in lines:
            low = ln.lower()
            if "руб" in low and not any(ch.isdigit() for ch in ln):
                words_line = ln.strip()
                break

    parts = [header]
    if customer_line:
        parts.append(customer_line)
    if inn_line:
        parts.append(inn_line)

    return "\n".join(parts).strip()


async def start_inn_search(message: _Message_inn, state: _FSMContext_inn):
    await message.answer("Готов к поиску… Пришлите ИНН и я покажу, что мне удалось найти.")
    await state.set_state("awaiting_inn_search")


async def handle_inn_input(message: _Message_inn, state: _FSMContext_inn):
    inn = message.text.strip()
    if not inn.isdigit():
        await message.answer("ИНН должен содержать только цифры. Попробуйте снова.")
        return

    await message.answer("Начал поиск, работаю с хранилищем…")
    results = []

    if not _os_inn.path.exists(GENERATED_PATH):
        await message.answer("❌ Папка с хранилищем не найдена.")
        await state.clear()
        return

    for root, _, files in _os_inn.walk(GENERATED_PATH):
        for file in files:
            if file.lower().endswith(".docx"):
                file_path = _os_inn.path.join(root, file)
                try:
                    doc = _Document_inn(file_path)
                    paragraphs = list(doc.paragraphs)
                    full_text = "\n".join(p.text for p in paragraphs)
                    if inn in full_text:
                        summary = build_inn_summary_from_paragraphs(paragraphs, file_path, inn)
                        results.append({"path": file_path, "summary": summary})
                except Exception:
                    continue

    if not results:
        await message.answer(f"Ничего не найдено по ИНН {inn} 😔")
        await state.clear()
        return

    await state.update_data(
        inn_search_results=results,
        inn_search_index=0,
        inn_search_inn=inn,
    )
    await state.set_state("inn_search_results")

    total = len(results)
    page_idx = 0
    page_text = f"{results[0]['summary']}\nСтраница {page_idx + 1}/{total}"
    await message.answer(page_text, reply_markup=inn_pagination_kb(), parse_mode=None)


async def inn_prev_page(callback: _CallbackQuery_inn, state: _FSMContext_inn):
    data = await state.get_data()
    results = data.get("inn_search_results") or []
    total = len(results)
    if total <= 1:
        await callback.answer()
        await callback.message.answer("Извини, в моей базе больше нет файлов с указанным ИНН")
        return

    idx = int(data.get("inn_search_index") or 0)
    if idx <= 0:
        await callback.answer()
        await callback.message.answer("Извини, в моей базе больше нет файлов с указанным ИНН")
        return

    idx -= 1
    await state.update_data(inn_search_index=idx)
    summary = results[idx].get("summary") or ""
    page_text = f"{summary}\nСтраница {idx + 1}/{total}"
    try:
        await callback.message.edit_text(page_text, reply_markup=inn_pagination_kb(), parse_mode=None)
    except Exception:
        await callback.message.answer(page_text, reply_markup=inn_pagination_kb(), parse_mode=None)
    await callback.answer()


async def inn_next_page(callback: _CallbackQuery_inn, state: _FSMContext_inn):
    data = await state.get_data()
    results = data.get("inn_search_results") or []
    total = len(results)
    if total <= 1:
        await callback.answer()
        await callback.message.answer("Извини, в моей базе больше нет файлов с указанным ИНН")
        return

    idx = int(data.get("inn_search_index") or 0)
    if idx >= total - 1:
        await callback.answer()
        await callback.message.answer("Извини, в моей базе больше нет файлов с указанным ИНН")
        return

    idx += 1
    await state.update_data(inn_search_index=idx)
    summary = results[idx].get("summary") or ""
    page_text = f"{summary}\nСтраница {idx + 1}/{total}"
    try:
        await callback.message.edit_text(page_text, reply_markup=inn_pagination_kb(), parse_mode=None)
    except Exception:
        await callback.message.answer(page_text, reply_markup=inn_pagination_kb(), parse_mode=None)
    await callback.answer()


async def inn_back_to_main(callback: _CallbackQuery_inn, state: _FSMContext_inn):
    await state.clear()
    await callback.answer()
    await callback.message.answer("Окей, возвращаю в главное меню 👇", reply_markup=main_kb())

# =====================================================

import os

if __name__ == "__main__":
    os.system("color 0A")  # 0 — черный фон, A — ярко-зелёный текст
    print("PROMO PRO АКТИВИРОВАН!Запуск всего самого важного и не только..")
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("\nОстановлено пользователем.")
    except Exception as e:
        logging.error("Фатальная ошибка запуска: %s", e)
        traceback.print_exc()
